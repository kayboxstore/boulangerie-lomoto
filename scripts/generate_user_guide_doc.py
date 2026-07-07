from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from boulangerie_app.version import APP_VERSION


OUTPUT = ROOT / "docs" / f"Guide-utilisateur-complet-Boulangerie-Lomoto-{APP_VERSION}.docx"
LOGO = ROOT / "boulangerie_app" / "assets" / "logo-boulangerie-lomoto.png"
SCREENSHOTS = ROOT / "presentations" / "guide-utilisateur" / "assets" / "screenshots"

RED = "B71924"
RED_DARK = "8B111A"
NAVY = "12233F"
BLUE = "1F4E78"
MUTED = "5A6570"
LIGHT_BLUE = "E8EEF5"
LIGHT_RED = "FDE8EA"
LIGHT_GOLD = "FFF4D8"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"
TEXT = "111827"
BORDER = "C8D4DF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for index, width in enumerate(widths):
        for cell in table.columns[index].cells:
            cell.width = Inches(width)


def set_table_borders(table, color: str = BORDER, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color: str = RED, size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)


def setup_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=3, width=Inches(6.8))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_width(table, [0.65, 4.95, 1.2])
    table.autofit = False
    left, center, right = table.rows[0].cells
    if LOGO.exists():
        left.paragraphs[0].add_run().add_picture(str(LOGO), width=Cm(1.1))
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    brand = center.paragraphs[0]
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(brand.add_run("BOULANGERIE LOMOTO"), size=11.5, color=RED, bold=True)
    sub = center.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(0)
    set_run_font(sub.add_run("Guide utilisateur complet"), size=8, color=MUTED, bold=True)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(rp.add_run(f"v{APP_VERSION}\n"), size=8.5, color=NAVY, bold=True)
    set_run_font(rp.add_run(date.today().strftime("%d/%m/%Y")), size=7.5, color=MUTED)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 10, 20, 10, 20)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rule = header.add_paragraph()
    rule.paragraph_format.space_after = Pt(0)
    paragraph_border_bottom(rule, RED, "8")

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run(f"© {date.today().year} Boulangerie Lomoto - General Investment Services (GIS). Tous droits réservés. | Page "), size=8, color=MUTED)
    add_page_number(paragraph)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(32)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(3.4))

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(4)
    set_run_font(kicker.add_run("MANUEL D'UTILISATION"), size=10, color=RED, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    set_run_font(title.add_run("Boulangerie Lomoto"), size=28, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(subtitle.add_run("Version Windows, Web Pro et Android"), size=14, color=BLUE, bold=True)

    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    table_width(meta, [1.9, 4.0])
    set_table_borders(meta, BORDER)
    rows = [
        ("Version application", APP_VERSION),
        ("Accès web", "https://app.boulangerie-lomoto.com"),
        ("APK Android", "BoulangerieLomoto-1.4.6-debug.apk pour recette; release signé après validation"),
        ("Public visé", "Admin, Directeur Général, Caissier, Production, Stock, Commandes"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_margins(cell, 100, 140, 100, 140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, bold=True, color=NAVY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    set_run_font(
        p.add_run(
            "Ce guide explique comment exploiter l'application au quotidien, gérer les accès, enregistrer les opérations, produire les rapports, clôturer les journées et travailler depuis Android."
        ),
        size=10.5,
        color=TEXT,
    )
    document.add_page_break()


def add_toc(document: Document) -> None:
    document.add_heading("Sommaire", level=1)
    entries = [
        "1. Vue d'ensemble",
        "2. Accès et installation",
        "3. Rôles et droits",
        "4. Connexion, session et sécurité du compte",
        "5. Tableau de bord",
        "6. Utilisateurs",
        "7. Stock",
        "8. Production",
        "9. Commandes",
        "10. Caisse",
        "11. Commissions",
        "12. Travailleurs et paies",
        "13. Rapports",
        "14. Clôture journalière",
        "15. Historique",
        "16. Sauvegardes et restauration",
        "17. Notifications e-mail",
        "18. Android",
        "19. Routine quotidienne",
        "20. Dépannage",
    ]
    for entry in entries:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(entry)
    document.add_page_break()


def add_callout(document: Document, title: str, body: str, fill: str = LIGHT_GOLD) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table_width(table, [6.5])
    set_table_borders(table, BORDER)
    set_table_row_pagination(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    set_run_font(p.add_run(f"{title} : "), size=10.5, color=NAVY, bold=True)
    set_run_font(p.add_run(body), size=10.5, color=TEXT)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def set_table_row_pagination(row, *, repeat_header: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if repeat_header and tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def compact_table_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05


def add_matrix(document: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_width(table, widths)
    set_table_borders(table, BORDER)
    set_table_row_pagination(table.rows[0], repeat_header=True)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade_cell(cell, header_fill)
        set_cell_margins(cell)
        compact_table_cell(cell)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        set_table_row_pagination(row)
        cells = row.cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            set_cell_margins(cells[index])
            compact_table_cell(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_screenshot(document: Document, filename: str, caption: str) -> None:
    path = SCREENSHOTS / filename
    if not path.exists():
        return
    with Image.open(path) as image:
        width, height = image.size
    max_width = Inches(6.2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=max_width)
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    set_run_font(caption_p.add_run(caption), size=8.5, color=MUTED, italic=True)


def add_section_overview(document: Document) -> None:
    document.add_heading("1. Vue d'ensemble", level=1)
    document.add_paragraph(
        "Boulangerie Lomoto est une application de gestion commerciale et opérationnelle synchronisée entre trois environnements : Windows, Web Pro et Android. "
        "Le PC serveur garde la base centrale, tandis que les postes Windows, les navigateurs et l'APK Android utilisent les mêmes données."
    )
    add_callout(
        document,
        "Principe central",
        "l'application ne doit jamais créer des données séparées par poste. Tout ce qui est enregistré sur Windows doit être visible sur le Web et Android, et inversement.",
        LIGHT_BLUE,
    )
    add_bullets(
        document,
        [
            "Windows sert aux postes de travail internes et au PC serveur.",
            "Web Pro sert aux accès distants depuis ordinateur, téléphone ou tablette.",
            "Android emballe la Web Pro dans une application installable.",
            "Cloudflare Tunnel publie le site sans ouvrir de port sur le routeur.",
            "Les rôles limitent l'accès : chacun voit ce qui concerne son travail.",
            "La clôture journalière fige les écritures de la journée et crée une trace exploitable.",
        ],
    )


def add_access_install(document: Document) -> None:
    document.add_heading("2. Accès et installation", level=1)
    document.add_heading("2.1 Accès Windows", level=2)
    add_numbered(
        document,
        [
            "Ouvrir l'application Windows depuis le raccourci Boulangerie Lomoto.",
            "Vérifier que le poste est en mode connecté au serveur central.",
            "Saisir l'identifiant ou l'e-mail, puis le mot de passe.",
            "Attendre l'ouverture du tableau de bord.",
        ],
    )
    document.add_heading("2.2 Accès Web Pro", level=2)
    add_numbered(
        document,
        [
            "Depuis le PC serveur : ouvrir http://127.0.0.1:8787.",
            "Depuis un téléphone sur le même Wi-Fi : utiliser l'adresse locale du PC serveur si connue.",
            "Depuis une autre connexion Internet : ouvrir https://app.boulangerie-lomoto.com.",
            "Se connecter avec le même compte que Windows.",
        ],
    )
    document.add_heading("2.3 Accès Android", level=2)
    add_numbered(
        document,
        [
            "Installer l'APK Android sur le téléphone.",
            "Ouvrir Boulangerie Lomoto depuis l'écran d'accueil Android.",
            "L'application charge automatiquement la Web Pro officielle.",
            "Se connecter avec son compte habituel.",
            "Tester l'accès en désactivant le Wi-Fi du serveur et en utilisant les données mobiles.",
        ],
    )
    add_callout(
        document,
        "Important Android",
        "la version debug sert à la recette. La livraison finale doit utiliser un APK release signé avec la clé Android conservée hors du PC serveur.",
        LIGHT_GOLD,
    )
    add_screenshot(document, "01-connexion.png", "Écran de connexion : l'utilisateur saisit son identifiant ou son e-mail, puis son mot de passe.")


def add_roles(document: Document) -> None:
    document.add_heading("3. Rôles et droits", level=1)
    add_matrix(
        document,
        ["Rôle", "Modules visibles", "Droit principal", "Limites"],
        [
            ["Admin", "Tous les modules", "Créer, modifier, supprimer, sauvegarder, restaurer, gérer les comptes", "Aucune limite métier; doit rester réservé au responsable"],
            ["Directeur Général", "Tous les modules", "Lecture globale, rapports, clôture journalière", "Lecture seule; ne modifie pas les écritures; ne rouvre pas une journée"],
            ["Caissier", "Caisse, Production, Commandes, Commissions, Travailleurs", "Gérer caisse, paies et certaines écritures autorisées", "Production, commandes et commissions peuvent être en lecture selon le contexte"],
            ["Chargé de la production", "Production", "Saisir la production de la journée", "Ne voit pas les modules hors production"],
            ["Gestionnaire de stock", "Stock", "Gérer approvisionnements, sorties et paramètres de stock", "Ne voit pas caisse, commandes, paies"],
            ["Gestionnaire des commandes", "Commandes, Commissions", "Enregistrer les commandes et consulter les commissions liées", "Ne voit pas stock ni caisse complète"],
        ],
        [1.25, 2.05, 1.65, 1.55],
    )
    document.add_heading("3.1 Règles de session", level=2)
    add_bullets(
        document,
        [
            "Un utilisateur ne doit avoir qu'une session active.",
            "Si l'application détecte une session déjà ouverte, elle propose de fermer l'ancienne session ou de rester sur l'ancienne.",
            "L'Admin voit les utilisateurs en ligne dans le module Utilisateurs.",
            "L'Admin peut forcer la déconnexion d'un utilisateur connecté.",
            "Une session inactive est nettoyée automatiquement afin d'éviter les faux blocages.",
        ],
    )


def add_login_security(document: Document) -> None:
    document.add_heading("4. Connexion, session et sécurité du compte", level=1)
    add_numbered(
        document,
        [
            "Saisir l'identifiant ou l'adresse e-mail.",
            "Saisir le mot de passe.",
            "Cocher Afficher le mot de passe uniquement si nécessaire.",
            "Cliquer sur Se connecter.",
            "Patienter pendant la barre de progression.",
        ],
    )
    add_bullets(
        document,
        [
            "Le mot de passe Admin et Directeur Général doit être fort : 14 caractères minimum avec majuscule, minuscule, chiffre et symbole.",
            "Les autres comptes doivent utiliser 12 caractères minimum.",
            "En cas de changement de mot de passe, le nouveau mot de passe est placé dans la file d'e-mail de l'utilisateur.",
            "Si un e-mail est absent lors de la création d'un compte, l'application peut générer une adresse du domaine de l'entreprise.",
            "Ne jamais partager un compte entre plusieurs personnes.",
        ],
    )
    add_callout(
        document,
        "Bon réflexe",
        "l'Admin doit désactiver ou changer immédiatement le mot de passe d'un compte lorsqu'une personne quitte l'entreprise.",
        LIGHT_RED,
    )


def add_dashboard(document: Document) -> None:
    document.add_heading("5. Tableau de bord", level=1)
    add_screenshot(document, "02-tableau-de-bord.png", "Tableau de bord : indicateurs mensuels et synthèse des opérations.")
    add_bullets(
        document,
        [
            "Les indicateurs clés sont mensuels et se réinitialisent au changement de mois.",
            "Les dettes non payées, commissions non payées et travailleurs non payés restent visibles jusqu'au règlement.",
            "Chaque rôle voit les indicateurs utiles à son travail.",
            "Le bouton Actualiser met à jour les données affichées.",
            "L'historique récent affiche les dernières actions utiles pour l'Admin.",
        ],
    )


def add_users(document: Document) -> None:
    document.add_heading("6. Utilisateurs", level=1)
    add_screenshot(document, "03-utilisateurs.png", "Module Utilisateurs : création, modification, rôle, e-mail et statut de connexion.")
    document.add_heading("6.1 Créer un utilisateur", level=2)
    add_numbered(
        document,
        [
            "Ouvrir Utilisateurs.",
            "Renseigner le nom complet.",
            "Saisir l'identifiant.",
            "Renseigner l'adresse e-mail.",
            "Choisir le rôle.",
            "Définir un mot de passe fort.",
            "Cliquer sur Enregistrer.",
        ],
    )
    document.add_heading("6.2 Modifier un utilisateur", level=2)
    add_numbered(
        document,
        [
            "Cliquer sur Charger ou sélectionner l'utilisateur.",
            "Vérifier les champs ramenés par l'application, y compris le mot de passe pour l'édition Admin.",
            "Modifier le rôle, l'e-mail ou le mot de passe si nécessaire.",
            "Cliquer sur Enregistrer.",
        ],
    )
    add_bullets(
        document,
        [
            "Un seul Directeur Général peut exister.",
            "Le dernier Admin ne doit pas être supprimé.",
            "Le statut En ligne/Hors ligne permet à l'Admin de suivre l'activité.",
            "Le bouton Déconnecter coupe la session active d'un utilisateur.",
            "La réinitialisation de la base est réservée à l'Admin et doit être utilisée seulement après sauvegarde.",
        ],
    )


def add_stock(document: Document) -> None:
    document.add_heading("7. Stock", level=1)
    add_screenshot(document, "04-stock.png", "Module Stock : synthèse, approvisionnements, sorties et journal.")
    add_screenshot(document, "05-approvisionnement-stock.png", "Approvisionnement : ajouter farine, levure, sel, huile et observations.")
    add_bullets(
        document,
        [
            "Le stock suit les entrées, les sorties et le stock restant.",
            "Les paramètres définissent les seuils d'alerte et les équivalences de consommation.",
            "Les approvisionnements ajoutent des quantités au stock.",
            "Les sorties déduisent les consommations utilisées par la production.",
            "Une date future est bloquée : on ne peut pas enregistrer demain si nous sommes aujourd'hui.",
            "Une journée clôturée ne peut plus être modifiée sauf réouverture Admin.",
        ],
    )
    document.add_heading("7.1 Procédure d'approvisionnement", level=2)
    add_numbered(
        document,
        [
            "Choisir la date du jour.",
            "Renseigner les quantités reçues.",
            "Ajouter une observation si nécessaire.",
            "Cliquer sur Enregistrer.",
            "Contrôler que la ligne apparaît dans l'historique.",
        ],
    )
    document.add_heading("7.2 Procédure de sortie", level=2)
    add_numbered(
        document,
        [
            "Choisir la date du jour.",
            "Renseigner les quantités utilisées.",
            "Cliquer sur Enregistrer la sortie.",
            "Vérifier le stock restant.",
        ],
    )


def add_production(document: Document) -> None:
    document.add_heading("8. Production", level=1)
    add_screenshot(document, "07-production.png", "Module Production : bacs commandés, livrés, donnés, restants, foutus et sacs utilisés.")
    add_bullets(
        document,
        [
            "Le Chargé de la production saisit la production de la journée.",
            "Le total produit est calculé à partir des bacs livrés, donnés, échantillons, restants et foutus.",
            "L'écart avec les commandes est calculé automatiquement.",
            "Le nombre de sacs utilisés permet de rapprocher production et stock.",
            "Les données de production alimentent les rapports.",
        ],
    )
    add_numbered(
        document,
        [
            "Ouvrir Production.",
            "Choisir la date du jour.",
            "Renseigner les bacs commandés.",
            "Renseigner les bacs livrés aux dépositaires et aux mamans.",
            "Renseigner les bacs donnés, échantillons, restants et foutus.",
            "Renseigner les sacs utilisés.",
            "Cliquer sur Enregistrer.",
        ],
    )


def add_orders(document: Document) -> None:
    document.add_heading("9. Commandes", level=1)
    add_screenshot(document, "06-commandes.png", "Module Commandes : client, statut, nombre de bacs, montant attendu, reçu, dette ou avance.")
    add_bullets(
        document,
        [
            "Les statuts de commande sont Dépositaire, Maman et Vente cash.",
            "Les tarifs sont calculés automatiquement : Dépositaire 4 100 FC, Maman 6 000 FC, Vente cash 4 350 FC.",
            "Le montant à percevoir se calcule à partir du nombre de bacs.",
            "Si le client paie moins, la différence devient une dette.",
            "Si le client paie plus, l'excédent est enregistré comme avance pour une prochaine commande.",
            "Le filtre permet d'afficher toutes les commandes, seulement les mamans ou seulement les dépositaires.",
            "Les dates futures sont bloquées.",
        ],
    )
    add_numbered(
        document,
        [
            "Ouvrir Commandes.",
            "Choisir la date du jour.",
            "Saisir le nom du client.",
            "Choisir le statut.",
            "Saisir le nombre de bacs.",
            "Vérifier le montant calculé.",
            "Saisir le montant reçu.",
            "Enregistrer et vérifier l'apparition dans l'historique.",
        ],
    )


def add_cash_commissions(document: Document) -> None:
    document.add_heading("10. Caisse", level=1)
    add_screenshot(document, "08-caisse.png", "Module Caisse : montants reçus, dettes payées, dépenses et solde.")
    add_bullets(
        document,
        [
            "La caisse rassemble les entrées, les dettes payées, les dépenses et le solde.",
            "Le montant reçu depuis les commandes est repris automatiquement.",
            "Les dépenses doivent être détaillées.",
            "Les dettes payées doivent être renseignées pour réduire les dettes accumulées.",
            "Le solde de caisse est calculé automatiquement.",
        ],
    )
    document.add_heading("11. Commissions", level=1)
    add_screenshot(document, "09-commissions.png", "Module Commissions : calcul et suivi des commissions liées aux commandes.")
    add_bullets(
        document,
        [
            "Les commissions sont calculées à partir des données de commandes.",
            "Le module affiche les commissions par date et par client concerné.",
            "Les commissions non payées restent visibles jusqu'au règlement.",
            "Le Gestionnaire des commandes peut voir ce qui concerne les commandes et commissions.",
        ],
    )


def add_workers(document: Document) -> None:
    document.add_heading("12. Travailleurs et paies", level=1)
    add_bullets(
        document,
        [
            "Le module Travailleurs conserve les informations des employés.",
            "L'e-mail du travailleur sert à recevoir les informations de paie.",
            "La date d'embauche permet de calculer l'ancienneté.",
            "L'ancienneté s'incrémente automatiquement chaque année.",
            "Les paies suivent les statuts : préparée, validée, payée.",
            "Les avances, retenues et primes ajustent le net à payer.",
            "Les travailleurs non payés restent visibles jusqu'au règlement.",
        ],
    )
    document.add_heading("12.1 Ajouter un travailleur", level=2)
    add_numbered(
        document,
        [
            "Ouvrir Travailleurs.",
            "Renseigner nom complet, fonction, téléphone, adresse et e-mail.",
            "Renseigner la date d'embauche.",
            "Définir le salaire mensuel.",
            "Enregistrer.",
        ],
    )
    document.add_heading("12.2 Enregistrer une paie", level=2)
    add_numbered(
        document,
        [
            "Sélectionner le travailleur.",
            "Choisir la période concernée.",
            "Renseigner brut, prime, avance et retenue.",
            "Vérifier le net calculé.",
            "Choisir le statut de paie.",
            "Enregistrer.",
        ],
    )


def add_reports_closure_history(document: Document) -> None:
    document.add_heading("13. Rapports", level=1)
    add_screenshot(document, "10-rapport-pdf.png", "Rapport PDF : génération depuis la date ou la période choisie.")
    add_screenshot(document, "11-rapport-excel.png", "Rapport Excel : journalier, mensuel, période, caisse hebdo ou bilan mensuel.")
    add_bullets(
        document,
        [
            "Le rapport journalier utilise uniquement la date de référence.",
            "Le rapport mensuel utilise le mois et l'année de la date de référence.",
            "Le rapport période utilise la date de début et la date de fin.",
            "Le bouton Actualiser recharge le contenu.",
            "Le bouton Afficher le dossier des rapports ouvre le dossier serveur des rapports.",
            "Après génération, le rapport doit s'ouvrir automatiquement.",
            "Le contenu du rapport dépend du rôle connecté.",
        ],
    )
    document.add_heading("14. Clôture journalière", level=1)
    add_bullets(
        document,
        [
            "La clôture fige les écritures de la journée.",
            "Elle génère un rapport signé et une sauvegarde.",
            "Le Directeur Général peut clôturer la journée.",
            "Seul l'Admin peut rouvrir une journée clôturée.",
            "La réouverture demande un motif.",
            "L'historique des clôtures conserve le jour, le statut, la date de clôture, la personne, le rôle, la réouverture et le motif.",
        ],
    )
    document.add_heading("15. Historique", level=1)
    add_bullets(
        document,
        [
            "L'historique affiche les 50 dernières actions par pages de 10.",
            "Les colonnes importantes sont date, utilisateur, rôle, module, action et détails.",
            "L'Admin peut filtrer par identifiant et par rôle.",
            "Les actions utiles sont journalisées : connexion, création, modification, suppression, rapports, clôtures, sauvegardes.",
        ],
    )


def add_backup_email_android(document: Document) -> None:
    document.add_heading("16. Sauvegardes et restauration", level=1)
    add_bullets(
        document,
        [
            "L'Admin peut sauvegarder la base depuis l'application.",
            "L'Admin peut restaurer une sauvegarde serveur.",
            "La restauration doit être précédée d'une sauvegarde de sécurité.",
            "Le bouton Voir les sauvegardes du serveur affiche les fichiers disponibles.",
            "La copie hebdomadaire hors PC serveur doit être contrôlée manuellement.",
        ],
    )
    document.add_heading("17. Notifications e-mail", level=1)
    add_bullets(
        document,
        [
            "Les e-mails sont placés dans une file d'attente pour ne pas bloquer l'enregistrement.",
            "La création utilisateur peut envoyer les identifiants.",
            "La modification de mot de passe peut envoyer le nouveau mot de passe.",
            "La paie peut notifier le travailleur selon son statut.",
            "Le bouton Relancer les envois force le traitement des e-mails en attente.",
            "Les e-mails légitimes dépendent aussi de SPF, DKIM, DMARC et de la réputation du domaine.",
        ],
    )
    document.add_heading("18. Android", level=1)
    add_bullets(
        document,
        [
            "L'APK Android ouvre la Web Pro officielle.",
            "Le package Android est com.kayboxstore.boulangerielomoto.",
            "La version actuelle est 1.4.6.",
            "Le trafic clair HTTP est désactivé : l'APK utilise HTTPS.",
            "La sauvegarde Android automatique est désactivée pour éviter l'exposition de session.",
            "Le téléphone doit avoir Internet pour utiliser l'application hors réseau local.",
        ],
    )
    add_matrix(
        document,
        ["Élément Android", "État"],
        [
            ["APK debug", "Généré pour recette : installer/output/android/1.4.6/BoulangerieLomoto-1.4.6-debug.apk"],
            ["APK release signé", "À générer après création de la clé Android"],
            ["Script de build", "scripts/build_android_apk.ps1"],
            ["Script installation USB", "scripts/install_android_debug.ps1"],
            ["Clé Android", "À créer avec scripts/create_android_keystore.ps1 et sauvegarder hors PC serveur"],
        ],
        [2.0, 4.5],
    )


def add_routine_troubleshooting(document: Document) -> None:
    document.add_heading("19. Routine quotidienne", level=1)
    add_matrix(
        document,
        ["Moment", "Action", "Responsable"],
        [
            ["Début de journée", "Vérifier serveur, connexion, stock initial et accès Web", "Admin / responsable"],
            ["Pendant la journée", "Saisir commandes, stock, production, caisse, paies selon rôle", "Chaque utilisateur"],
            ["Après chaque saisie", "Vérifier l'historique et les tableaux", "Utilisateur concerné"],
            ["Fin de journée", "Contrôler les chiffres, générer rapport, clôturer", "Admin ou DG"],
            ["Chaque semaine", "Contrôler la sauvegarde externe", "Admin"],
            ["Chaque mois", "Contrôler indicateurs mensuels, dettes, paies, commissions", "Admin / DG"],
        ],
        [1.35, 3.8, 1.35],
    )
    document.add_heading("20. Dépannage", level=1)
    add_matrix(
        document,
        ["Symptôme", "Cause probable", "Solution"],
        [
            ["Le téléphone ne charge pas 127.0.0.1", "127.0.0.1 désigne le téléphone lui-même", "Utiliser l'adresse locale du PC serveur ou le domaine public"],
            ["Le site distant ne s'ouvre pas", "Tunnel Cloudflare arrêté, PC serveur éteint ou Internet coupé", "Vérifier services Windows, Cloudflared et connexion"],
            ["Session déjà ouverte", "Session active ailleurs ou ancienne session non nettoyée", "Choisir fermer l'ancienne session ou demander à l'Admin"],
            ["Impossible d'enregistrer", "Journée clôturée, date future ou rôle non autorisé", "Vérifier date, rôle et statut de clôture"],
            ["Rapport ne s'ouvre pas", "Chemin de rapport indisponible ou droits Windows", "Utiliser le bouton Afficher le dossier des rapports"],
            ["E-mail non reçu", "File d'attente, service e-mail ou filtrage destinataire", "Relancer les envois et vérifier configuration e-mail"],
            ["Données non visibles", "Page non actualisée ou filtre actif", "Cliquer Actualiser et vérifier date/filtre"],
        ],
        [1.7, 2.2, 2.6],
    )


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    setup_styles(document)
    add_header_footer(document)
    add_cover(document)
    add_toc(document)
    add_section_overview(document)
    add_access_install(document)
    add_roles(document)
    add_login_security(document)
    add_dashboard(document)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_users(document)
    add_stock(document)
    add_production(document)
    add_orders(document)
    add_cash_commissions(document)
    add_workers(document)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_reports_closure_history(document)
    add_backup_email_android(document)
    add_routine_troubleshooting(document)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
