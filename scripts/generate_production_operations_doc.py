from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from boulangerie_app.version import APP_VERSION  # noqa: E402


OUTPUT = ROOT / "docs" / f"Guide-production-securite-exploitation-Boulangerie-Lomoto-{APP_VERSION}.docx"
LOGO = ROOT / "boulangerie_app" / "assets" / "logo-boulangerie-lomoto.png"

RED = RGBColor(183, 25, 36)
NAVY = RGBColor(16, 24, 39)
BLUE = RGBColor(22, 58, 99)
MUTED = RGBColor(82, 94, 111)
LIGHT_BLUE = "EAF0F7"
LIGHT_RED = "FDE8EA"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in (
        ("Heading 1", 16, RED),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    table = header.add_table(rows=1, cols=3, width=Inches(7.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [2.0, 11.0, 4.0])
    left, center, right = table.rows[0].cells
    if LOGO.exists():
        left.paragraphs[0].add_run().add_picture(str(LOGO), width=Cm(1.45))
    title = center.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BOULANGERIE LOMOTO")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RED
    subtitle = center.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(0)
    run = subtitle.add_run("GUIDE PRODUCTION, SECURITE ET EXPLOITATION")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = BLUE
    meta = right.paragraphs[0]
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run(f"Version {APP_VERSION}\n")
    run.bold = True
    run.font.color.rgb = NAVY
    run = meta.add_run(date.today().strftime("%d/%m/%Y"))
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    line = header.add_paragraph()
    p_pr = line._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B71924")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"© {date.today().year} Boulangerie Lomoto - General Investment Services (GIS). "
        "Solution initiee par Augustin Kayembe. Tous droits reserves."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("Guide final de production et d'exploitation")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_after = Pt(14)
    run = intro.add_run(
        "Ce document regroupe les actions a appliquer pour exploiter Boulangerie Lomoto "
        "en production : sauvegardes, surveillance, securite, e-mails, APK Android, "
        "rapports, maintenance Admin et controles quotidiens."
    )
    run.font.size = Pt(11.5)
    run.font.color.rgb = MUTED

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, [5.2, 5.2, 5.2])
    values = [
        ("Plateformes", "Windows, Web, APK Android"),
        ("Serveur", "PC local connecte H24"),
        ("Acces public", "Cloudflare Tunnel, sans port routeur"),
    ]
    for index, (label, value) in enumerate(values):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        r.bold = True
        r.font.color.rgb = BLUE
        r = p.add_run(value)
        r.font.color.rgb = NAVY


def h(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def p(document: Document, text: str, bold_prefix: str = "") -> None:
    paragraph = document.add_paragraph()
    if bold_prefix:
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)


def bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def checklist_table(document: Document, rows: list[tuple[str, str, str]]) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, [4.0, 7.2, 4.6])
    for idx, label in enumerate(("Controle", "Action", "Frequence / responsable")):
        cell = table.rows[0].cells[idx]
        cell.text = label
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = NAVY
    for controle, action, frequence in rows:
        cells = table.add_row().cells
        for idx, value in enumerate((controle, action, frequence)):
            cells[idx].text = value
            set_cell_margins(cells[idx])


def command_block(document: Document, commands: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, [16.0])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    for idx, command in enumerate(commands):
        if idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY


def add_sections(document: Document) -> None:
    h(document, "1. Sauvegardes hors PC serveur")
    p(document, "La base reste sur le PC serveur. La sauvegarde locale ne suffit pas : il faut une copie externe utilisable meme si le PC tombe en panne.")
    checklist_table(
        document,
        [
            ("Sauvegarde quotidienne", "Tache planifiee vers l'API locale de maintenance.", "Chaque soir, PC serveur"),
            ("Sauvegarde externe", "Disque USB nomme LOMOTO_BACKUP ou destination manuelle.", "Chaque dimanche"),
            ("Test restauration", "Verifier l'integrite SQLite de la derniere sauvegarde.", "Chaque semaine, Admin"),
            ("Retention", "Conserver au moins 12 sauvegardes externes.", "Automatique"),
        ],
    )
    command_block(
        document,
        [
            r"powershell -ExecutionPolicy Bypass -File .\scripts\installer-taches-production-lomoto.ps1",
            r"powershell -ExecutionPolicy Bypass -File .\scripts\tester-restauration-sauvegarde-lomoto.ps1",
        ],
    )

    h(document, "2. Surveillance du serveur et du tunnel")
    bullets(
        document,
        [
            "Le service Windows BoulangerieLomotoCentralServer doit rester en cours d'execution.",
            "Cloudflare Tunnel reste l'unique ouverture publique : aucune redirection de port routeur.",
            "La tache de surveillance verifie le service local, relance Cloudflare Tunnel si necessaire et traite la file e-mail.",
        ],
    )
    checklist_table(
        document,
        [
            ("Local", "http://127.0.0.1:8787/api/health doit repondre ok.", "Automatique + controle Admin"),
            ("Public", "https://boulangerie-lomoto.com/api/health doit repondre ok.", "Automatique + telephone hors Wi-Fi"),
            ("Logs", r"C:\ProgramData\BoulangerieLomoto\maintenance", "A consulter en cas de panne"),
        ],
    )

    h(document, "3. Securite renforcee")
    checklist_table(
        document,
        [
            ("2FA Cloudflare", "Activer une application d'authentification sur le compte Cloudflare.", "Obligatoire"),
            ("2FA e-mail", "Activer la double authentification sur l'adresse e-mail principale.", "Obligatoire"),
            ("Mots de passe", "Admin et DG : 14 caracteres minimum, non reutilises.", "A chaque creation"),
            ("Sessions", "Une seule session active par utilisateur; Admin peut deconnecter un compte.", "Automatique"),
            ("Tentatives", "Limiter les echecs de connexion et journaliser les refus.", "Automatique"),
            ("Historique", "Archiver avant effacement; ne pas supprimer sans raison.", "Admin"),
        ],
    )

    h(document, "4. APK Android final")
    bullets(
        document,
        [
            "L'APK embarque la version web officielle : aucune logique metier dupliquee.",
            "Le logo, le nom, le splash screen et le mode plein ecran sont prepares par le script de build.",
            "Pour Play Store, il faut une cle de signature conservee hors PC serveur.",
        ],
    )
    command_block(
        document,
        [
            r"powershell -ExecutionPolicy Bypass -File .\scripts\create_android_keystore.ps1",
            r"powershell -ExecutionPolicy Bypass -File .\scripts\build_android_apk.ps1 -Release",
        ],
    )

    h(document, "5. E-mails transactionnels")
    p(document, "Les e-mails sont mis en file d'attente pour eviter de ralentir les enregistrements. La surveillance traite ensuite les messages en arriere-plan.")
    checklist_table(
        document,
        [
            ("SPF/DKIM/DMARC", "Configurer les entrees DNS du domaine pour authentifier les messages.", "Avant production mail"),
            ("Adresse d'envoi", "Utiliser notifications@boulangerie-lomoto.com.", "Admin"),
            ("Test d'envoi", "Envoyer un mail de test depuis Utilisateurs > Notifications.", "Apres configuration"),
            ("File d'attente", "Relancer les envois si un message reste en echec.", "Admin"),
        ],
    )

    h(document, "6. Rapports")
    bullets(
        document,
        [
            "Journalier : une date de reference.",
            "Mensuel : le mois de la date de reference.",
            "Periode : date de debut et date de fin.",
            "Les avances clients sont affichees a part et ne gonflent pas le solde caisse.",
            "Apres generation, le PDF peut s'ouvrir automatiquement sur le PC serveur.",
        ],
    )

    h(document, "7. Interface Web, mobile et tablette")
    checklist_table(
        document,
        [
            ("PC", "Tester Chrome/Edge avec cache vide apres mise a jour.", "A chaque version"),
            ("Telephone", "Tester hors Wi-Fi serveur via reseau mobile.", "Avant validation domaine"),
            ("APK", "Tester saisie formulaire, clavier, deconnexion 15 minutes.", "Avant livraison"),
            ("Tableaux", "Verifier les lignes larges et les boutons sur petit ecran.", "Recette mobile"),
        ],
    )

    h(document, "8. Administration et maintenance")
    bullets(
        document,
        [
            "Le module Utilisateurs affiche en ligne / hors ligne, plateforme, IP et derniere activite.",
            "L'Admin peut deconnecter un utilisateur actif.",
            "L'Admin peut sauvegarder, restaurer et reinitialiser la base depuis Windows.",
            "L'ecran Etat systeme expose chemins, sauvegardes, e-mails, sessions et licence.",
        ],
    )

    h(document, "9. Documentation et recette")
    checklist_table(
        document,
        [
            ("Guide Admin", "Conserver ce guide avec les guides par role.", "Derniere version"),
            ("Recette complete", r"python scripts\run_lomoto_recette_complete.py", "Avant livraison"),
            ("Installateur", "Compiler l'installateur apres validation recette.", "Avant remise client"),
            ("Support", "Noter toute intervention dans l'historique ou un journal externe.", "Continu"),
        ],
    )

    h(document, "Procedure de controle quotidien", 2)
    bullets(
        document,
        [
            "Verifier que le domaine s'ouvre depuis un telephone en 4G.",
            "Verifier que les utilisateurs connectes sont normaux.",
            "Verifier qu'il n'y a pas d'e-mails en echec.",
            "Verifier qu'une sauvegarde recente existe.",
            "Verifier que le disque externe est branche le jour de sauvegarde hebdomadaire.",
        ],
    )


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    setup_document(document)
    add_header_footer(document)
    add_cover(document)
    add_sections(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
