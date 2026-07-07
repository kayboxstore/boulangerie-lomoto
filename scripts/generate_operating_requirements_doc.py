from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from boulangerie_app.version import APP_VERSION

LOGO = ROOT / "boulangerie_app" / "assets" / "logo-boulangerie-lomoto.png"
OUTPUT = ROOT / "docs" / "Besoins-techniques-et-mise-en-service-Boulangerie-Lomoto.docx"
PRICE_OUTPUT = ROOT / "docs" / "Materiels-outils-prix-Boulangerie-Lomoto.docx"
USD_TO_CDF = 2300

RED = "B71924"
BLUE = "163A63"
NAVY = "0F2347"
LIGHT_BLUE = "EAF0F7"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"
TEXT = RGBColor(20, 23, 31)
MUTED = RGBColor(90, 101, 112)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_header(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=3, width=Inches(6.8))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(2.2)
    table.columns[1].width = Cm(11.6)
    table.columns[2].width = Cm(3.2)
    left, center, right = table.rows[0].cells
    if LOGO.exists():
        left.paragraphs[0].add_run().add_picture(str(LOGO), width=Cm(1.55))
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    brand = center.paragraphs[0]
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("BOULANGERIE LOMOTO")
    run.bold = True
    run.font.name = "Poppins"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(RED)
    subtitle = center.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(0)
    run = subtitle.add_run("PLAN DE MISE EN SERVICE ET BESOINS TECHNIQUES")
    run.bold = True
    run.font.name = "Poppins"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right_p.add_run(f"Version {APP_VERSION}\n")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run = right_p.add_run(date.today().strftime("%d/%m/%Y"))
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 25, 40, 25, 40)

    border = header.add_paragraph()
    border.paragraph_format.space_after = Pt(0)
    p_pr = border._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RED)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    run = paragraph.add_run(
        f"© {date.today().year} Boulangerie Lomoto - General Investment Services (GIS). "
        "Tous droits réservés."
    )
    run.font.name = "Poppins"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("Conditions pour le bon fonctionnement de l'application")
    run.bold = True
    run.font.name = "Poppins"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.space_after = Pt(12)
    run = intro.add_run(
        "Architecture retenue : un PC serveur local reste allumé 24 h/24 et conserve la base centrale. "
        "Les applications Windows, Web, téléphone, tablette et future APK Android utilisent cette même base. "
        "Les accès extérieurs passent par Cloudflare Tunnel et le domaine boulangerie-lomoto.com, sans redirection de port routeur."
    )
    run.font.color.rgb = TEXT


def heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Poppins"
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = RGBColor.from_string(RED if level == 1 else BLUE)


def bullet(document: Document, text: str, bold_prefix: str = "") -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)


def add_requirements_table(document: Document) -> None:
    rows = [
        ("PC serveur", "1", "Indispensable", "Windows 10/11 64 bits, Intel Core i5/Ryzen 5 ou mieux, 8 Go RAM minimum, SSD 256 Go minimum."),
        ("Connexion Ethernet", "1", "Indispensable", "Relier le PC serveur au routeur par câble réseau; éviter le Wi-Fi pour le serveur."),
        ("Onduleur PC", "1", "Indispensable", "650 à 1000 VA minimum pour protéger le PC et permettre un arrêt propre."),
        ("Mini UPS routeur", "1", "Indispensable", "Tension et connecteurs compatibles avec le routeur/ONT; autonomie visée de 4 à 8 heures."),
        ("Disque USB externe", "1", "Indispensable", "SSD/HDD 500 Go ou plus, nommé LOMOTO_BACKUP, pour une copie hebdomadaire hors PC serveur."),
        ("Routeur stable", "1", "Indispensable", "Réservation d'adresse IP locale pour le PC serveur et Wi-Fi correct pour téléphones/tablettes."),
        ("Connexion Internet", "1", "Indispensable", "Fibre, 4G ou autre accès stable avec débit montant correct; IP publique fixe non obligatoire avec Cloudflare Tunnel."),
        ("Connexion de secours", "1", "Recommandé", "Routeur 4G/5G ou second opérateur pour réduire les interruptions."),
        ("Ventilation et emplacement", "1", "Recommandé", "PC serveur dans un endroit sec, ventilé, sécurisé et inaccessible au public."),
        ("Téléphone/tablette de test", "1", "Recommandé", "Android récent pour la recette responsive, PWA et APK."),
    ]
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Élément", "Qté", "Priorité", "Caractéristiques / rôle")
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = label
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for item, quantity, priority, detail in rows:
        cells = table.add_row().cells
        values = (item, quantity, priority, detail)
        for index, value in enumerate(values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])
        if priority == "Indispensable":
            set_cell_shading(cells[2], "FDE8EA")
        else:
            set_cell_shading(cells[2], LIGHT_BLUE)


def format_usd_range(low: int, high: int) -> str:
    if low == high:
        return f"{low} USD"
    return f"{low} - {high} USD"


def format_cdf_range(low: int, high: int) -> str:
    low_cdf = low * USD_TO_CDF
    high_cdf = high * USD_TO_CDF
    if low == high:
        return f"{low_cdf:,} FC".replace(",", " ")
    return f"{low_cdf:,} - {high_cdf:,} FC".replace(",", " ")


def add_price_proposal_table(document: Document) -> None:
    rows = [
        ("PC serveur reconditionne Core i5/Ryzen 5, 8 Go RAM, SSD 256/512 Go", "Indispensable si aucun PC fiable", 350, 650, "Choisir une machine sobre, stable et facile a reparer."),
        ("Upgrade RAM ou SSD du PC serveur existant", "Recommande si le PC est lent", 35, 120, "A faire seulement si le PC actuel est conserve."),
        ("Onduleur PC 650 a 1000 VA", "Indispensable", 80, 160, "Protege le PC serveur et laisse le temps d'arreter proprement."),
        ("Mini UPS routeur/ONT", "Indispensable", 25, 60, "Garde Internet actif pendant une coupure courte."),
        ("Disque USB externe 1 To", "Indispensable", 55, 100, "Sauvegarde hebdomadaire hors PC serveur."),
        ("Cable Ethernet, multiprise, rangement", "Indispensable", 15, 40, "Connexion serveur stable et installation propre."),
        ("Routeur 4G/5G ou connexion de secours", "Recommande", 70, 180, "Reduit les coupures d'acces distant."),
        ("Forfait Internet principal", "Mensuel", 25, 80, "A dimensionner selon l'operateur disponible."),
        ("Domaine boulangerie-lomoto.com", "Annuel", 10, 25, "Deja choisi; renouvellement annuel a surveiller."),
        ("Cloudflare Tunnel / DNS / HTTPS", "Service", 0, 0, "Plan gratuit suffisant pour l'architecture retenue."),
        ("Envoi e-mails transactionnels", "Service", 0, 20, "Cloudflare Email Sending teste; prevoir une alternative SMTP si les quotas changent."),
        ("Compte Google Play Console", "Une seule fois", 25, 25, "Necessaire pour publier l'APK sur Play Store."),
        ("Telephone Android de test", "Recommande", 120, 300, "Un telephone recent suffit pour la recette mobile."),
    ]
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Element / outil", "Priorite", "Prix USD", "Prix FC approx.", "Conseil")
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = label
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for item, priority, low, high, detail in rows:
        cells = table.add_row().cells
        values = (item, priority, format_usd_range(low, high), format_cdf_range(low, high), detail)
        for index, value in enumerate(values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])
        if "Indispensable" in priority:
            set_cell_shading(cells[1], "FDE8EA")
        elif "Recommande" in priority:
            set_cell_shading(cells[1], LIGHT_BLUE)
        else:
            set_cell_shading(cells[1], LIGHT_GRAY)


def add_budget_summary_table(document: Document) -> None:
    rows = [
        ("Budget minimum avec PC serveur deja disponible", 210, 460, "Onduleur, mini UPS, disque externe, cables, Internet et domaine."),
        ("Budget minimum si achat d'un PC serveur", 560, 1110, "PC serveur fiable + equipements indispensables."),
        ("Budget confort avec secours Internet et telephone de test", 750, 1590, "Ajoute routeur/connexion de secours et appareil de recette."),
        ("Couts annuels/mensuels a suivre", 35, 125, "Domaine annuel, Internet mensuel, eventuel SMTP/Email selon quotas."),
    ]
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Scenario", "Budget USD", "Budget FC approx.", "Contenu")
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = label
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for scenario, low, high, detail in rows:
        cells = table.add_row().cells
        for index, value in enumerate((scenario, format_usd_range(low, high), format_cdf_range(low, high), detail)):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])


def add_checklist_table(document: Document) -> None:
    items = [
        ("PC serveur connecté par câble et allumé 24 h/24", "À valider"),
        ("Onduleur PC et mini UPS routeur testés pendant une coupure", "À valider"),
        ("Adresse IP locale du serveur réservée dans le routeur", "À valider"),
        ("Pare-feu Windows sans ouverture publique du port Web 8787", "Validé côté PC"),
        ("Cloudflare Tunnel utilisé sans redirection de port routeur", "Tunnel actif; routeur à contrôler"),
        ("Sauvegarde automatique locale et copie hebdomadaire sur disque USB", "Tâche utilisateur installée"),
        ("Restauration d'une sauvegarde testée", "À valider"),
        ("Rôles et droits vérifiés avec chaque profil", "À valider"),
        ("Blocage des dates futures vérifié sur tous les modules", "À valider"),
        ("Synchronisation Windows/Web vérifiée dans les deux sens", "À valider"),
        ("Affichage testé sur PC, téléphone et tablette", "À valider"),
        ("Test distant effectué depuis une autre connexion Internet", "À confirmer depuis téléphone hors Wi-Fi"),
        ("Nom de domaine, DNS, certificat HTTPS et tunnel configurés", "Actif sur boulangerie-lomoto.com"),
        ("2FA Cloudflare et adresse e-mail du compte activées", "Action manuelle propriétaire"),
        ("APK Android signé, installé et testé", "Après recette mobile finale"),
        ("Test continu de 72 heures sans perte de données", "À valider"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, label in enumerate(("Contrôle avant mise en production", "État")):
        cell = table.rows[0].cells[index]
        cell.text = label
        set_cell_shading(cell, RED)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for control, status in items:
        cells = table.add_row().cells
        cells[0].text = control
        cells[1].text = status
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
        set_cell_shading(cells[1], LIGHT_GRAY if status == "À valider" else LIGHT_BLUE)


def add_security_status_table(document: Document) -> None:
    rows = [
        (
            "Accès public",
            "Cloudflare Tunnel uniquement",
            "Le port Web 8787 reste local au PC serveur. Ne pas créer de redirection NAT/port forwarding sur le routeur.",
        ),
        (
            "Cloudflare",
            "2FA à activer manuellement",
            "Activer l'authentification à deux facteurs, sauvegarder les codes de récupération hors du PC serveur.",
        ),
        (
            "Adresse e-mail du compte",
            "2FA à activer manuellement",
            "Activer la double authentification de l'adresse utilisée pour Cloudflare et le registrar du domaine.",
        ),
        (
            "Mots de passe Admin/DG",
            "Durcis dans l'application",
            "14 caractères minimum, majuscule, minuscule, chiffre et symbole; éviter nom, identifiant et mots évidents.",
        ),
        (
            "Mots de passe autres rôles",
            "Durcis dans l'application",
            "12 caractères minimum, majuscule, minuscule, chiffre et symbole.",
        ),
        (
            "Sauvegarde hors PC serveur",
            "Automatisée chaque semaine",
            "Tâche Windows hebdomadaire utilisateur vers disque USB externe LOMOTO_BACKUP, avec conservation des 12 dernières copies. À passer en mode administrateur/SYSTEM si la session Windows doit rester fermée.",
        ),
        (
            "Test distant",
            "À valider sur téléphone",
            "Ouvrir https://app.boulangerie-lomoto.com depuis les données mobiles, Wi-Fi serveur désactivé.",
        ),
    ]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Mesure", "État", "Procédure / remarque")
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = label
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for measure, status, detail in rows:
        cells = table.add_row().cells
        for index, value in enumerate((measure, status, detail)):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])
        set_cell_shading(cells[1], LIGHT_BLUE if "activer manuellement" not in status.lower() else "FDE8EA")


def build_document() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    add_header(section)
    add_footer(section)

    styles = document.styles
    styles["Normal"].font.name = "Poppins"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = TEXT
    styles["List Bullet"].font.name = "Poppins"
    styles["List Bullet"].font.size = Pt(9.5)

    add_title(document)

    heading(document, "1. Matériels nécessaires")
    add_requirements_table(document)

    heading(document, "2. Propositions de prix")
    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(6)
    note.add_run(
        f"Base de conversion indicative utilisee dans ce document : 1 USD = {USD_TO_CDF:,} FC. "
        "Les prix sont des fourchettes d'achat a confirmer chez les fournisseurs locaux avant depense."
    )
    add_price_proposal_table(document)
    heading(document, "3. Budgets conseilles", level=2)
    add_budget_summary_table(document)

    heading(document, "4. Logiciels, services et outils")
    bullet(document, "Application Windows Boulangerie Lomoto installée sur le PC serveur.", "Application Windows : ")
    bullet(document, "service Web Pro démarré automatiquement sur le PC serveur.", "Application Web : ")
    bullet(document, "choix obligatoire entre PC serveur principal et poste client pendant l'installation.", "Mode d'installation : ")
    bullet(document, "Cloudflare Tunnel utilisé pour publier le Web en HTTPS sans redirection de port routeur.", "Accès Internet : ")
    bullet(document, "boulangerie-lomoto.com configuré dans Cloudflare; renouvellement annuel à suivre.", "Domaine : ")
    bullet(document, "Cloudflare Email Sending configuré; le test d'envoi a été reçu. Continuer la recette avec les vrais e-mails transactionnels.", "E-mails : ")
    bullet(document, "Chrome, Edge, Firefox ou Safari à jour sur les appareils clients.", "Navigateurs : ")
    bullet(document, "Android Studio, SDK Android, Java et clé de signature conservée en lieu sûr; APK Android signé en version 1.4.6.", "APK Android : ")
    bullet(document, "antivirus Windows actif avec exclusions limitées aux services de l'application si nécessaire.", "Protection : ")

    heading(document, "5. Architecture réseau retenue")
    bullet(document, "Le PC serveur conserve l'unique base de données de production.")
    bullet(document, "Sur le même réseau, le téléphone ouvre l'adresse locale du PC, par exemple http://192.168.1.225:8787.")
    bullet(document, "Depuis une autre connexion Internet, le Web utilise https://app.boulangerie-lomoto.com via Cloudflare Tunnel.")
    bullet(document, "Le routeur reçoit une réservation DHCP pour que l'adresse locale du PC serveur ne change pas.")
    bullet(document, "Aucune ouverture directe du fichier SQLite, aucune exposition publique du port Web et aucune redirection de port routeur.")
    bullet(document, "Le mini UPS maintient le routeur et l'ONT; l'onduleur classique protège le PC serveur.")

    heading(document, "6. Synchronisation Windows, Web et mobile")
    bullet(document, "Toutes les écritures passent par le serveur central et la même base.")
    bullet(document, "Une commande saisie sur Windows devient visible sur le Web après actualisation automatique ou manuelle.")
    bullet(document, "Une opération saisie sur le Web est immédiatement disponible pour l'application Windows.")
    bullet(document, "L'application ne doit jamais basculer silencieusement vers une base locale séparée lorsque le serveur central est indisponible.")
    bullet(document, "Le PC serveur doit rester allumé et connecté pour les accès distants dans l'architecture locale retenue.")

    document.add_page_break()
    heading(document, "7. Comptes, sécurité et connexion rapide")
    add_security_status_table(document)
    bullet(document, "Le premier administrateur est créé uniquement sur le PC serveur principal et uniquement si la base centrale ne contient encore aucun utilisateur.")
    bullet(document, "Une réinstallation ou l'installation d'un poste client conserve les comptes existants et ne redemande jamais cette configuration.")
    bullet(document, "Chaque utilisateur reçoit son propre rôle et ne voit que les modules autorisés.")
    bullet(document, "L'adresse e-mail est obligatoire pour les nouveaux utilisateurs et peut servir à la connexion.")
    bullet(document, "Après création d'un compte, ses informations de connexion sont placées dans la file d'envoi e-mail.")
    bullet(document, "Les travailleurs reçoivent par e-mail les changements de statut de leur paie : préparée, validée ou payée.")
    bullet(document, "Le navigateur peut mémoriser la connexion avec son gestionnaire de mots de passe.")
    bullet(document, "L'application ne conserve plus une case de mémorisation visible; les champs de connexion sont vidés après une connexion réussie.")
    bullet(document, "Le domaine est actif et les paramètres Cloudflare Email Sending ont été testés avec succès.")
    bullet(document, "Le mot de passe communiqué par e-mail doit être changé par l'utilisateur dès sa première connexion.")
    bullet(document, "Activer l'authentification à deux facteurs sur Cloudflare, le compte du domaine et l'adresse e-mail utilisée.")

    heading(document, "8. Sauvegardes et continuité")
    bullet(document, "Sauvegarde automatique quotidienne de la base sur le PC serveur.")
    bullet(document, "Copie hebdomadaire automatisée sur un disque USB externe nommé LOMOTO_BACKUP, puis disque débranché après contrôle.")
    bullet(document, "Conservation de plusieurs sauvegardes, pas seulement la dernière.")
    bullet(document, "Test de restauration au moins une fois par mois.")
    bullet(document, "Sauvegarde créée automatiquement lors de la clôture journalière.")
    bullet(document, "Après hébergement, ajouter une copie chiffrée hors site.")

    heading(document, "9. Téléphones, tablettes et Android")
    bullet(document, "L'interface Web est responsive et utilisable sur téléphone, tablette et ordinateur.")
    bullet(document, "La PWA pourra être installée depuis le navigateur une fois le site servi en HTTPS.")
    bullet(document, "L'APK Android ouvre l'URL HTTPS finale, utilise le logo Boulangerie Lomoto et respecte l'affichage mobile.")
    bullet(document, "La clé de signature Android doit être sauvegardée; sans elle, les futures mises à jour de l'APK deviennent impossibles.")

    heading(document, "10. Exploitation quotidienne")
    bullet(document, "Vérifier chaque matin que le PC serveur, le service Web et la connexion Internet sont actifs.")
    bullet(document, "Clôturer la journée; seul l'administrateur peut la rouvrir.")
    bullet(document, "Ne jamais enregistrer une opération pour une date future.")
    bullet(document, "Contrôler les alertes de stock, dettes, commissions et paies non réglées.")
    bullet(document, "Installer les mises à jour d'abord sur le serveur, puis vérifier le Web et les postes Windows.")

    document.add_section(WD_SECTION.NEW_PAGE)
    heading(document, "11. Recette et contrôles avant exploitation finale")
    add_checklist_table(document)

    heading(document, "12. Suite logique recommandée")
    steps = [
        "Vérifier le serveur Web local, le service Windows et le tunnel Cloudflare après redémarrage du PC serveur.",
        "Exécuter une recette complète Windows/Web pendant 72 heures avec la vraie base.",
        "Tester une coupure de courant, le mini UPS, l'onduleur et la reprise automatique des services.",
        "Tester une sauvegarde puis une restauration sur une copie de la base.",
        "Valider tous les rôles, les dates, les rapports et la synchronisation dans les deux sens.",
        "Activer la 2FA Cloudflare, la 2FA de l'adresse e-mail et conserver les codes de récupération hors PC serveur.",
        "Tester https://app.boulangerie-lomoto.com depuis un téléphone en données mobiles, Wi-Fi du serveur désactivé.",
        "Vérifier dans le routeur qu'aucune redirection de port vers le PC serveur n'existe.",
        "Installer la dernière APK signée sur le téléphone et vérifier l'icône, la barre système, les thèmes et les formulaires.",
    ]
    for index, step in enumerate(steps, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"{index}. ")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(RED)
        paragraph.add_run(step)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    set_cell = document.add_table(rows=1, cols=1)
    set_cell.style = "Table Grid"
    cell = set_cell.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    run = p.add_run("Décision recommandée : ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    p.add_run(
        "ne pas annoncer l'application comme définitivement livrée avant la validation 2FA, sauvegarde/restauration, "
        "test téléphone hors Wi-Fi serveur et stabilité du PC serveur pendant 72 heures."
    )

    document.save(OUTPUT)
    document.save(PRICE_OUTPUT)
    print(PRICE_OUTPUT)


if __name__ == "__main__":
    build_document()
