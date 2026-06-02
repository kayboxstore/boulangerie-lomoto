from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
SUPPORT_DIR = ROOT / "docs" / "supports" / "1.3.12"
GUIDE_MD = ROOT / "docs" / "guide-exploitation-1.3.12.md"
RECETTE_MD = ROOT / "docs" / "fiche-recette-1.3.12.md"
GUIDE_DOCX = SUPPORT_DIR / "Guide-exploitation-Boulangerie-Lomoto-1.3.12.docx"
GUIDE_PDF = SUPPORT_DIR / "Guide-exploitation-Boulangerie-Lomoto-1.3.12.pdf"
RECETTE_XLSX = SUPPORT_DIR / "Fiche-recette-Boulangerie-Lomoto-1.3.12.xlsx"
README_SUPPORTS = SUPPORT_DIR / "README.md"
LOGO = ROOT / "boulangerie_app" / "assets" / "logo-boulangerie-lomoto.png"
P_REGULAR = ROOT / "boulangerie_app" / "fonts" / "Poppins-Regular.ttf"
P_BOLD = ROOT / "boulangerie_app" / "fonts" / "Poppins-Bold.ttf"

VERSION = "1.3.12"
TODAY = date(2026, 6, 2)
BLUE = "1F4E78"
NAVY = "0B2545"
RED = "B71C1C"
LIGHT_BLUE = "D9EAF7"
LIGHT = "F7FAFC"
GREEN = "1E7D32"
AMBER = "F4B400"
TEXT = "1F2933"


GUIDE_SECTIONS = [
    (
        "Objectif du guide",
        [
            "Ce guide accompagne la version 1.3.12 de Boulangerie Lomoto. Il sert à installer, exploiter, tester et maintenir l’application dans un environnement réel.",
            "Il doit être gardé avec les installateurs officiels afin que l’administrateur, le caissier et les gestionnaires puissent suivre les mêmes consignes.",
        ],
    ),
    (
        "Livrables de la version 1.3.12",
        [
            "Setup officiel : installer/output/1.3.12/BoulangerieLomotoSetup.exe",
            "Setup démo : installer/output/1.3.12-demo/BoulangerieLomotoDemoSetup.exe",
            "Release GitHub : https://github.com/kayboxstore/boulangerie-lomoto/releases/tag/v1.3.12",
            "Manifeste de mise à jour : https://raw.githubusercontent.com/kayboxstore/boulangerie-lomoto-updates/main/update.json",
        ],
    ),
    (
        "Installation du poste serveur / admin",
        [
            "Installer la version officielle sur le PC principal.",
            "Lancer l’application en administrateur si Windows le demande, puis se connecter avec le compte administrateur.",
            "Ouvrir Paramètres réseau et installer ou mettre à jour le service central.",
            "Vérifier que le service central démarre sans erreur et que les ports TCP 8765 et UDP 8766 sont autorisés dans le pare-feu.",
            "Générer un rapport journalier de test pour vérifier les polices, le logo et les tableaux.",
        ],
    ),
    (
        "Installation d’un poste client",
        [
            "Installer le setup officiel sur le poste client.",
            "Ouvrir l’application et se connecter avec le compte prévu pour l’utilisateur.",
            "Vérifier que l’application reste en mode connecté après la connexion.",
            "Tester l’ouverture des modules autorisés selon le rôle de l’utilisateur.",
            "Si le mode connecté ne s’active pas, vérifier que le PC serveur est allumé, que le service central tourne et que les deux postes sont sur le même réseau local.",
        ],
    ),
    (
        "Rôles et accès",
        [
            "Administrateur : accès complet, utilisateurs, sauvegarde, restauration, rapports complets, travailleurs et paies.",
            "Caissier : commandes visibles, caisse, commissions et argent; lecture seule sur les éléments de commandes si nécessaire.",
            "Gestionnaire des commandes : commandes, commissions et production.",
            "Gestionnaire de stock : stock, approvisionnement et rapports de stock.",
            "Chaque utilisateur ne doit voir que les boutons utiles à son rôle.",
        ],
    ),
    (
        "Routine quotidienne recommandée",
        [
            "Avant l’ouverture : démarrer le PC serveur et vérifier que le service central fonctionne.",
            "Pendant la journée : saisir les commandes, la production, le stock, la caisse et les paiements de dettes au bon moment.",
            "Avant la clôture : vérifier les dettes, les dépenses, les sacs utilisés, les commissions et les alertes critiques.",
            "Après la clôture : générer le rapport journalier et conserver une sauvegarde.",
        ],
    ),
    (
        "Rapports 1.3.12",
        [
            "Le premier tableau récapitulatif affiche maintenant le Net à payer des commissions et le Solde après paiement des commissions.",
            "Le rapport mensuel ne reprend plus la liste complète de toutes les commandes du mois.",
            "Le rapport mensuel ne reprend plus la liste des personnes qui ont payé leurs dettes.",
            "Les commandes mensuelles sont synthétisées par statut pour garder le rapport propre et lisible.",
            "Les tableaux PDF et Excel utilisent des en-têtes bleu foncé, des lignes alternées et une meilleure hiérarchie visuelle.",
        ],
    ),
    (
        "Sécurité et bonnes pratiques",
        [
            "Changer le mot de passe administrateur avant une utilisation réelle chez un client.",
            "Ne pas partager le compte admin avec les utilisateurs simples.",
            "Faire une sauvegarde avant toute mise à jour importante.",
            "Tester la version démo séparément pour les présentations commerciales.",
            "Ne jamais supprimer le dossier de données sans sauvegarde préalable.",
        ],
    ),
    (
        "Procédure de mise à jour",
        [
            "L’application vérifie automatiquement les mises à jour selon le manifeste en ligne.",
            "Si une mise à jour est proposée, fermer les fenêtres ouvertes et accepter l’installation.",
            "Après installation, rouvrir l’application et vérifier que la version affichée est bien 1.3.12.",
            "Sur le poste serveur, mettre à jour le service central si l’application le demande.",
            "Sur les postes clients, tester la connexion au serveur après la mise à jour.",
        ],
    ),
    (
        "Dépannage rapide",
        [
            "Mode local après connexion : vérifier le service central, le pare-feu et la détection automatique du serveur.",
            "Boutons grisés : vérifier le rôle connecté et l’état de clôture de la journée.",
            "Rapport incomplet : vérifier que les données du module concerné ont bien été saisies.",
            "Setup bloqué sur un fichier service : arrêter le service central, fermer l’application puis relancer le setup.",
            "Erreur de mot de passe : attendre la fin du blocage temporaire si plusieurs tentatives incorrectes ont été faites.",
        ],
    ),
]


TEST_CASES = [
    ("Installation", "Installer la version officielle sur le PC serveur", "Admin", "Lancer le setup 1.3.12 puis ouvrir l’application.", "L’application s’ouvre en version 1.3.12.", "Haute"),
    ("Installation", "Installer la version démo sur un PC séparé", "Admin", "Lancer le setup démo et se connecter avec l’utilisateur démo.", "La démo s’ouvre sans mélanger les données officielles.", "Moyenne"),
    ("Mise à jour", "Détection de la version 1.3.12", "Admin", "Ouvrir l’application depuis une version précédente.", "La mise à jour 1.3.12 est proposée.", "Haute"),
    ("Mode connecté", "Démarrer le service central", "Admin", "Paramètres réseau > Installer/mettre à jour le service puis démarrer.", "Le service central démarre sans erreur.", "Haute"),
    ("Mode connecté", "Connexion d’un poste client", "Utilisateur", "Se connecter depuis un autre poste du même réseau.", "Le mode connecté reste actif après connexion.", "Haute"),
    ("Mode connecté", "Détection automatique du serveur", "Utilisateur", "Ouvrir l’application client sans saisir manuellement l’adresse.", "L’adresse du serveur est récupérée automatiquement.", "Haute"),
    ("Connexion", "Connexion administrateur", "Admin", "Saisir l’identifiant admin et le mot de passe.", "Le tableau de bord admin s’affiche.", "Haute"),
    ("Connexion", "Blocage après mauvais mot de passe", "Tous", "Saisir 5 mauvais mots de passe.", "Le compte est bloqué temporairement.", "Haute"),
    ("Tableau de bord", "Affichage selon rôle admin", "Admin", "Se connecter admin.", "Tous les modules autorisés sont visibles.", "Haute"),
    ("Tableau de bord", "Affichage selon rôle caissier", "Caissier", "Se connecter caissier.", "Seuls les boutons utiles au caissier sont visibles.", "Haute"),
    ("Commandes", "Commande Maman avec dette", "Gestionnaire commandes", "Saisir une commande Maman avec montant reçu inférieur au montant attendu.", "La dette est calculée correctement.", "Haute"),
    ("Commandes", "Blocage montant reçu supérieur", "Gestionnaire commandes", "Saisir un montant reçu supérieur au montant à percevoir.", "L’enregistrement est refusé.", "Haute"),
    ("Commandes", "Détection client similaire", "Gestionnaire commandes", "Saisir Akonga après Akonga Mapasa le même jour.", "L’application propose modification ou client différent.", "Moyenne"),
    ("Commandes", "Grilles par statut", "Gestionnaire commandes", "Ouvrir les commandes après plusieurs statuts.", "Les dépositaires et mamans/vente cash sont lisibles séparément ou filtrables.", "Moyenne"),
    ("Commissions", "Synchronisation automatique", "Gestionnaire commandes", "Enregistrer une commande Maman.", "La commission apparaît sans bouton Enregistrer.", "Haute"),
    ("Commissions", "Dépositaire sans commission", "Gestionnaire commandes", "Enregistrer un dépositaire.", "La commission reste à zéro.", "Haute"),
    ("Caisse", "Paiement de dette", "Caissier", "Saisir une dette payée aujourd’hui avec nom et montant.", "Le total des entrées augmente et la liste est visible dans la grille.", "Haute"),
    ("Caisse", "Aucune dette à payer", "Caissier", "Ouvrir une journée sans dette accumulée.", "Le message indique que personne n’a payé car il n’y a pas de dette accumulée.", "Moyenne"),
    ("Caisse", "Rapport mensuel allégé", "Caissier/Admin", "Générer le rapport mensuel.", "La liste des payeurs de dettes n’est pas reprise dans le mensuel.", "Haute"),
    ("Stock", "Approvisionnement", "Gestionnaire stock", "Ajouter un approvisionnement.", "Le stock augmente et l’historique est visible.", "Haute"),
    ("Stock", "Cohérence sacs utilisés", "Gestionnaire stock", "Saisir une valeur différente de la production.", "L’application affiche une alerte de non-correspondance.", "Haute"),
    ("Production", "Saisie manuelle sacs utilisés", "Gestionnaire commandes", "Saisir la production et le nombre de sacs utilisés.", "La valeur saisie est conservée et reportée.", "Haute"),
    ("Travailleurs", "Créer un travailleur", "Admin", "Ajouter un travailleur actif.", "Le travailleur apparaît dans la grille.", "Moyenne"),
    ("Travailleurs", "Créer une paie", "Admin", "Ajouter une paie avec prime/avance/retenue.", "Le net à payer est calculé correctement.", "Moyenne"),
    ("Travailleurs", "Bouton Fermer", "Admin", "Ouvrir Travailleurs puis cliquer Fermer.", "La fenêtre se ferme proprement.", "Moyenne"),
    ("Rapports PDF", "Rapport journalier complet", "Admin", "Générer le rapport journalier.", "Logo, en-tête, tableaux et solde après commissions sont visibles.", "Haute"),
    ("Rapports PDF", "Rapport mensuel commandes", "Admin", "Générer le rapport mensuel avec plusieurs commandes.", "Le rapport affiche une synthèse par statut, pas la liste complète.", "Haute"),
    ("Rapports PDF", "Rapport par rôle stock", "Gestionnaire stock", "Générer un rapport avec compte stock.", "Seules les données stock sont visibles.", "Moyenne"),
    ("Rapports Excel", "Export journalier Excel", "Admin", "Générer l’export Excel journalier.", "Les onglets utiles existent et les montants sont formatés.", "Moyenne"),
    ("Rapports Excel", "Export mensuel Excel", "Admin", "Générer l’export Excel mensuel.", "La feuille commandes est synthétique.", "Haute"),
    ("Sauvegarde", "Sauvegarde manuelle", "Admin", "Lancer une sauvegarde.", "Un fichier de sauvegarde est créé.", "Haute"),
    ("Restauration", "Restauration contrôlée", "Admin", "Restaurer une sauvegarde de test.", "Les données restaurées sont correctes.", "Haute"),
    ("Clôture", "Clôturer journée", "Admin/Caissier", "Clôturer après saisie complète.", "La journée est verrouillée contre les modifications non autorisées.", "Haute"),
    ("Sécurité", "Lecture seule caissier sur commandes", "Caissier", "Ouvrir commandes avec compte caissier.", "Le caissier voit sans modifier si prévu.", "Moyenne"),
    ("Performance", "Ouverture modules", "Tous", "Cliquer sur les grands modules.", "Le délai d’ouverture reste très court.", "Moyenne"),
]


def set_run_font(run, *, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill=LIGHT_BLUE, header_text=NAVY) -> None:
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, name="Calibri", size=9.5, color=TEXT)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, name="Calibri", size=9.5, color=header_text, bold=True)


def add_doc_paragraph(doc: Document, text: str, style: str | None = None, *, bold=False, color=TEXT) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, name="Calibri", size=11, color=color, bold=bold)


def add_doc_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        set_run_font(run, name="Calibri", size=10.5, color=TEXT)


def build_guide_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if style_name == "Heading 2" else 18)
        style.paragraph_format.space_after = Pt(7 if style_name == "Heading 2" else 10)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"Boulangerie Lomoto | Guide exploitation {VERSION}")
    set_run_font(run, name="Calibri", size=9, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Document interne - Kay Box Store")
    set_run_font(run, name="Calibri", size=9, color="6B7280")

    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(str(LOGO), width=Inches(1.2))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Guide d’exploitation")
    set_run_font(run, name="Calibri", size=24, color=RED, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(f"Boulangerie Lomoto - Version {VERSION}")
    set_run_font(run, name="Calibri", size=13, color=NAVY, bold=True)

    meta = doc.add_table(rows=5, cols=2)
    meta.rows[0].cells[0].text = "Version"
    meta.rows[0].cells[1].text = VERSION
    meta.rows[1].cells[0].text = "Date"
    meta.rows[1].cells[1].text = TODAY.strftime("%d/%m/%Y")
    meta.rows[2].cells[0].text = "Public"
    meta.rows[2].cells[1].text = "Administrateur, caissier, gestionnaires et support technique"
    meta.rows[3].cells[0].text = "But"
    meta.rows[3].cells[1].text = "Installer, exploiter, tester et maintenir l’application"
    meta.rows[4].cells[0].text = "Point d’attention"
    meta.rows[4].cells[1].text = "Changer le mot de passe admin avant utilisation réelle"
    style_table(meta)

    add_doc_paragraph(
        doc,
        "À retenir : le PC serveur doit rester allumé pour que les postes clients travaillent en mode connecté. "
        "Les mises à jour et sauvegardes doivent être vérifiées avant toute exploitation chez un client.",
        bold=True,
        color=NAVY,
    )

    for title_text, items in GUIDE_SECTIONS:
        doc.add_heading(title_text, level=1)
        if len(items) <= 2:
            for item in items:
                add_doc_paragraph(doc, item)
        else:
            add_doc_bullets(doc, items)

    doc.add_heading("Checklist de démarrage réel", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, header_text in enumerate(["Point", "Responsable", "Résultat attendu", "Validé"]):
        table.rows[0].cells[idx].text = header_text
    rows = [
        ("Installer le setup serveur", "Admin", "Application ouverte en 1.3.12", "☐"),
        ("Mettre à jour le service central", "Admin", "Service démarré sans erreur", "☐"),
        ("Installer un poste client", "Support", "Connexion automatique au serveur", "☐"),
        ("Tester les rôles", "Admin", "Boutons visibles selon rôle", "☐"),
        ("Générer rapports PDF/Excel", "Admin/Caissier", "Rapports lisibles et complets", "☐"),
        ("Créer sauvegarde", "Admin", "Fichier de sauvegarde disponible", "☐"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
    style_table(table)

    doc.save(GUIDE_DOCX)


def pdf_style_sheet():
    if P_REGULAR.exists() and P_BOLD.exists():
        pdfmetrics.registerFont(TTFont("Poppins", str(P_REGULAR)))
        pdfmetrics.registerFont(TTFont("Poppins-Bold", str(P_BOLD)))
        regular, bold = "Poppins", "Poppins-Bold"
    else:
        regular, bold = "Helvetica", "Helvetica-Bold"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("GuideTitle", parent=styles["Title"], fontName=bold, fontSize=24, textColor=colors.HexColor(f"#{RED}"), alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle("GuideSubtitle", parent=styles["Normal"], fontName=bold, fontSize=13, textColor=colors.HexColor(f"#{NAVY}"), alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle("GuideH1", parent=styles["Heading1"], fontName=bold, fontSize=14, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle("GuideBody", parent=styles["Normal"], fontName=regular, fontSize=9.5, leading=12.5, textColor=colors.HexColor(f"#{TEXT}"), alignment=TA_LEFT, spaceAfter=5))
    styles.add(ParagraphStyle("GuideBullet", parent=styles["Normal"], fontName=regular, fontSize=9.2, leading=12, leftIndent=12, firstLineIndent=-6, textColor=colors.HexColor(f"#{TEXT}"), spaceAfter=3))
    styles.add(ParagraphStyle("GuideNote", parent=styles["Normal"], fontName=bold, fontSize=9.5, leading=12.5, textColor=colors.HexColor(f"#{NAVY}"), backColor=colors.HexColor(f"#{LIGHT_BLUE}"), borderPadding=8, spaceAfter=8))
    return styles, regular, bold


def build_guide_pdf() -> None:
    styles, _regular, bold = pdf_style_sheet()
    doc = SimpleDocTemplate(
        str(GUIDE_PDF),
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=f"Guide exploitation Boulangerie Lomoto {VERSION}",
    )
    story = []
    if LOGO.exists():
        logo = Image(str(LOGO), width=24 * mm, height=24 * mm)
        logo.hAlign = "CENTER"
        story.append(logo)
    story.append(Paragraph("Guide d’exploitation", styles["GuideTitle"]))
    story.append(Paragraph(f"Boulangerie Lomoto - Version {VERSION}", styles["GuideSubtitle"]))
    story.append(
        Table(
            [
                ["Version", VERSION, "Date", TODAY.strftime("%d/%m/%Y")],
                ["Public", "Admin, caissier, gestionnaires", "But", "Exploitation et recette"],
            ],
            colWidths=[26 * mm, 58 * mm, 24 * mm, 58 * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT_BLUE}")),
                    ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor(f"#{LIGHT}")),
                    ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(f"#{NAVY}")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CAD6E2")),
                    ("FONTNAME", (0, 0), (-1, -1), "Poppins"),
                    ("FONTNAME", (0, 0), (0, -1), bold),
                    ("FONTNAME", (2, 0), (2, -1), bold),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            ),
        )
    )
    story.append(Spacer(1, 6 * mm))
    story.append(
        Paragraph(
            "À retenir : le PC serveur doit rester allumé pour que les postes clients travaillent en mode connecté. "
            "Avant une exploitation réelle, vérifier l’installation, le service central, les rôles, les rapports et la sauvegarde.",
            styles["GuideNote"],
        )
    )
    for section_title, items in GUIDE_SECTIONS:
        block = [Paragraph(section_title, styles["GuideH1"])]
        if len(items) <= 2:
            block.extend(Paragraph(item, styles["GuideBody"]) for item in items)
        else:
            block.extend(Paragraph(f"• {item}", styles["GuideBullet"]) for item in items)
        story.append(KeepTogether(block))
    story.append(Paragraph("Checklist de démarrage réel", styles["GuideH1"]))
    table_rows = [["Point", "Responsable", "Résultat attendu", "Validé"]]
    table_rows.extend(
        [
            ["Installer le setup serveur", "Admin", "Application ouverte en 1.3.12", "☐"],
            ["Mettre à jour le service central", "Admin", "Service démarré sans erreur", "☐"],
            ["Installer un poste client", "Support", "Connexion automatique au serveur", "☐"],
            ["Tester les rôles", "Admin", "Boutons visibles selon rôle", "☐"],
            ["Générer rapports PDF/Excel", "Admin/Caissier", "Rapports lisibles et complets", "☐"],
            ["Créer sauvegarde", "Admin", "Fichier de sauvegarde disponible", "☐"],
        ]
    )
    story.append(
        Table(
            [[Paragraph(str(cell), styles["GuideBody"]) for cell in row] for row in table_rows],
            colWidths=[48 * mm, 34 * mm, 72 * mm, 18 * mm],
            repeatRows=1,
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), bold),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{LIGHT}")]),
                    ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(f"#{NAVY}")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CAD6E2")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            ),
        )
    )
    doc.build(story)


def build_markdown_sources() -> None:
    guide_lines = [
        f"# Guide d’exploitation - Boulangerie Lomoto {VERSION}",
        "",
        f"Date : {TODAY.strftime('%d/%m/%Y')}",
        "",
        "Ce guide accompagne l’installation et l’exploitation de la version 1.3.12.",
        "",
    ]
    for title_text, items in GUIDE_SECTIONS:
        guide_lines.extend([f"## {title_text}", ""])
        if len(items) <= 2:
            for item in items:
                guide_lines.extend([item, ""])
        else:
            guide_lines.extend([f"- {item}" for item in items])
            guide_lines.append("")
    GUIDE_MD.write_text("\n".join(guide_lines), encoding="utf-8")

    recette_lines = [
        f"# Fiche de recette utilisateur - Boulangerie Lomoto {VERSION}",
        "",
        "Objectif : vérifier que la version 1.3.12 est prête pour une utilisation réelle.",
        "",
        "| Domaine | Scénario | Rôle | Résultat attendu | Priorité |",
        "|---|---|---|---|---|",
    ]
    for domain, scenario, role, _steps, expected, priority in TEST_CASES:
        recette_lines.append(f"| {domain} | {scenario} | {role} | {expected} | {priority} |")
    RECETTE_MD.write_text("\n".join(recette_lines), encoding="utf-8")


def build_recette_xlsx() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Recette complète"
    summary = wb.create_sheet("Résumé")
    instructions = wb.create_sheet("Mode d'emploi")

    title_fill = PatternFill("solid", fgColor=NAVY)
    header_fill = PatternFill("solid", fgColor=BLUE)
    section_fill = PatternFill("solid", fgColor=LIGHT_BLUE)
    light_fill = PatternFill("solid", fgColor=LIGHT)
    pass_fill = PatternFill("solid", fgColor="E7F4EA")
    fail_fill = PatternFill("solid", fgColor="FCE8E6")
    border = Border(
        left=Side(style="thin", color="CAD6E2"),
        right=Side(style="thin", color="CAD6E2"),
        top=Side(style="thin", color="CAD6E2"),
        bottom=Side(style="thin", color="CAD6E2"),
    )

    for sheet in (ws, summary, instructions):
        sheet.sheet_view.showGridLines = False

    ws.merge_cells("A1:J1")
    ws["A1"] = f"Fiche de recette utilisateur - Boulangerie Lomoto {VERSION}"
    ws["A1"].fill = title_fill
    ws["A1"].font = Font(name="Poppins", size=16, bold=True, color="FFFFFF")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28
    ws.merge_cells("A2:J2")
    ws["A2"] = "Cochez chaque scénario après test réel sur le serveur et au moins un poste client."
    ws["A2"].fill = section_fill
    ws["A2"].font = Font(name="Poppins", size=11, bold=True, color=NAVY)
    ws["A2"].alignment = Alignment(horizontal="left", vertical="center")

    headers = ["N°", "Domaine", "Scénario", "Rôle", "Étapes", "Résultat attendu", "Priorité", "Statut", "Commentaires", "Date/Testeur"]
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(4, col, header)
        cell.fill = header_fill
        cell.font = Font(name="Poppins", size=10, bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    for row_index, case in enumerate(TEST_CASES, start=5):
        domain, scenario, role, steps, expected, priority = case
        values = [row_index - 4, domain, scenario, role, steps, expected, priority, "À tester", "", ""]
        for col, value in enumerate(values, start=1):
            cell = ws.cell(row_index, col, value)
            cell.font = Font(name="Poppins", size=10, color=TEXT)
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            cell.border = border
            if row_index % 2 == 0:
                cell.fill = light_fill
        ws.cell(row_index, 1).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row_index, 7).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row_index, 8).alignment = Alignment(horizontal="center", vertical="center")
    ws.freeze_panes = "A5"
    ws.auto_filter.ref = f"A4:J{len(TEST_CASES) + 4}"
    widths = [6, 18, 34, 18, 48, 46, 12, 14, 34, 18]
    for col_index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(col_index)].width = width
    for row_idx in range(5, len(TEST_CASES) + 5):
        ws.row_dimensions[row_idx].height = 48
    validation = DataValidation(type="list", formula1='"À tester,Réussi,Échec,À reprendre"', allow_blank=False)
    ws.add_data_validation(validation)
    validation.add(f"H5:H{len(TEST_CASES) + 4}")

    from openpyxl.formatting.rule import CellIsRule

    ws.conditional_formatting.add(
        f"H5:H{len(TEST_CASES) + 4}",
        CellIsRule(operator="equal", formula=['"Réussi"'], fill=pass_fill),
    )
    ws.conditional_formatting.add(
        f"H5:H{len(TEST_CASES) + 4}",
        CellIsRule(operator="equal", formula=['"Échec"'], fill=fail_fill),
    )
    ws.conditional_formatting.add(
        f"H5:H{len(TEST_CASES) + 4}",
        CellIsRule(operator="equal", formula=['"À reprendre"'], fill=PatternFill("solid", fgColor="FFF4CE")),
    )

    summary.merge_cells("A1:F1")
    summary["A1"] = f"Résumé de recette - Version {VERSION}"
    summary["A1"].fill = title_fill
    summary["A1"].font = Font(name="Poppins", size=16, bold=True, color="FFFFFF")
    summary["A1"].alignment = Alignment(horizontal="center")
    summary_rows = [
        ("Version", VERSION),
        ("Date de préparation", TODAY.strftime("%d/%m/%Y")),
        ("Tests prévus", f"=COUNTA('Recette complète'!A5:A{len(TEST_CASES)+4})"),
        ("Réussis", f'=COUNTIF(\'Recette complète\'!H5:H{len(TEST_CASES)+4},"Réussi")'),
        ("Échecs", f'=COUNTIF(\'Recette complète\'!H5:H{len(TEST_CASES)+4},"Échec")'),
        ("À reprendre", f'=COUNTIF(\'Recette complète\'!H5:H{len(TEST_CASES)+4},"À reprendre")'),
        ("À tester", f'=COUNTIF(\'Recette complète\'!H5:H{len(TEST_CASES)+4},"À tester")'),
        ("Décision", '=IF(E5>0,"Bloquer la livraison",IF(E6>0,"Corriger avant client",IF(E7=0,"Prêt pour client","Recette incomplète")))'),
    ]
    for row_index, (label, value) in enumerate(summary_rows, start=3):
        summary.cell(row_index, 1, label)
        summary.cell(row_index, 2, value)
        for col in range(1, 3):
            cell = summary.cell(row_index, col)
            cell.border = border
            cell.font = Font(name="Poppins", size=11, bold=(col == 1), color=TEXT)
            cell.alignment = Alignment(horizontal="left", vertical="center")
            if col == 1:
                cell.fill = section_fill
    summary["D3"] = "Statut"
    summary["E3"] = "Nombre"
    for cell_ref in ("D3", "E3"):
        summary[cell_ref].fill = header_fill
        summary[cell_ref].font = Font(name="Poppins", size=10, bold=True, color="FFFFFF")
        summary[cell_ref].alignment = Alignment(horizontal="center")
        summary[cell_ref].border = border
    status_rows = [("Réussi", "=B6"), ("Échec", "=B7"), ("À reprendre", "=B8"), ("À tester", "=B9")]
    for idx, (status, formula) in enumerate(status_rows, start=4):
        summary.cell(idx, 4, status)
        summary.cell(idx, 5, formula)
        for col in (4, 5):
            cell = summary.cell(idx, col)
            cell.border = border
            cell.font = Font(name="Poppins", size=10, color=TEXT)
            cell.alignment = Alignment(horizontal="center")
            if idx % 2:
                cell.fill = light_fill
    chart = BarChart()
    chart.title = "État de la recette"
    chart.y_axis.title = "Nombre de tests"
    chart.x_axis.title = "Statut"
    data = Reference(summary, min_col=5, min_row=3, max_row=7)
    cats = Reference(summary, min_col=4, min_row=4, max_row=7)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 7
    chart.width = 12
    summary.add_chart(chart, "D10")
    summary.column_dimensions["A"].width = 24
    summary.column_dimensions["B"].width = 24
    summary.column_dimensions["D"].width = 18
    summary.column_dimensions["E"].width = 14

    instructions.merge_cells("A1:F1")
    instructions["A1"] = "Mode d’emploi de la fiche de recette"
    instructions["A1"].fill = title_fill
    instructions["A1"].font = Font(name="Poppins", size=16, bold=True, color="FFFFFF")
    instructions["A1"].alignment = Alignment(horizontal="center")
    instruction_rows = [
        ("1", "Installer ou mettre à jour le poste serveur avec le setup officiel 1.3.12."),
        ("2", "Installer au moins un poste client et vérifier le mode connecté."),
        ("3", "Exécuter chaque scénario de l’onglet Recette complète."),
        ("4", "Mettre le statut Réussi, Échec ou À reprendre après chaque test."),
        ("5", "Corriger tous les échecs avant de livrer chez un client."),
        ("6", "Faire une sauvegarde après validation complète."),
    ]
    for row_index, (step, text) in enumerate(instruction_rows, start=3):
        instructions.cell(row_index, 1, step)
        instructions.cell(row_index, 2, text)
        for col in (1, 2):
            cell = instructions.cell(row_index, col)
            cell.border = border
            cell.font = Font(name="Poppins", size=11, bold=(col == 1), color=TEXT)
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            if row_index % 2 == 0:
                cell.fill = light_fill
    instructions.column_dimensions["A"].width = 8
    instructions.column_dimensions["B"].width = 90
    for row_idx in range(3, 9):
        instructions.row_dimensions[row_idx].height = 28

    wb.save(RECETTE_XLSX)
    loaded = load_workbook(RECETTE_XLSX, data_only=False)
    assert "Recette complète" in loaded.sheetnames
    assert "Résumé" in loaded.sheetnames
    assert len(TEST_CASES) == loaded["Recette complète"].max_row - 4


def build_readme() -> None:
    README_SUPPORTS.write_text(
        "\n".join(
            [
                f"# Supports version {VERSION}",
                "",
                "Ce dossier contient les supports de validation et d’exploitation de la version 1.3.12.",
                "",
                "- `Guide-exploitation-Boulangerie-Lomoto-1.3.12.docx` : guide Word à remettre à l’administrateur.",
                "- `Guide-exploitation-Boulangerie-Lomoto-1.3.12.pdf` : version PDF du guide.",
                "- `Fiche-recette-Boulangerie-Lomoto-1.3.12.xlsx` : checklist Excel pour tester l’application avant livraison.",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    SUPPORT_DIR.mkdir(parents=True, exist_ok=True)
    build_markdown_sources()
    build_guide_docx()
    build_guide_pdf()
    build_recette_xlsx()
    build_readme()
    print("Supports générés :")
    print(GUIDE_DOCX)
    print(GUIDE_PDF)
    print(RECETTE_XLSX)


if __name__ == "__main__":
    main()
