from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt

from generate_user_guide_doc import (
    APP_VERSION,
    BORDER,
    LIGHT_BLUE,
    LIGHT_GOLD,
    LIGHT_RED,
    LOGO,
    MUTED,
    NAVY,
    RED,
    ROOT,
    add_bullets,
    add_callout,
    add_matrix,
    add_numbered,
    add_page_number,
    add_screenshot,
    set_cell_margins,
    set_run_font,
    set_table_borders,
    setup_styles,
    shade_cell,
    table_width,
)


OUTPUT_DIR = ROOT / "docs" / "guides-utilisateur-par-module"


GUIDES = [
    {
        "number": "01",
        "slug": "connexion-et-compte",
        "title": "Connexion et sécurité du compte",
        "roles": "Tous les utilisateurs",
        "objective": "Se connecter correctement, gérer une session unique, changer son mot de passe et réagir à un conflit de session.",
        "screenshots": [("01-connexion.png", "Écran de connexion de l'application.")],
        "fields": [
            ["Identifiant ou e-mail", "Compte attribué par l'Admin. Ne pas partager cet identifiant."],
            ["Mot de passe", "Secret personnel. Utiliser Afficher le mot de passe seulement pour vérifier une saisie."],
            ["Se connecter", "Valide les informations et ouvre le tableau de bord autorisé par le rôle."],
            ["Changer le mot de passe", "Obligatoire lorsque le compte utilise encore son mot de passe initial."],
        ],
        "procedures": [
            ("Connexion normale", [
                "Ouvrir Windows, la Web Pro ou l'APK Android.",
                "Saisir l'identifiant ou l'adresse e-mail.",
                "Saisir le mot de passe.",
                "Cliquer sur Se connecter et attendre la fin de la progression.",
                "Vérifier le nom et le rôle affichés dans l'application.",
            ]),
            ("Session déjà ouverte", [
                "Lire la plateforme et l'appareil indiqués dans le message.",
                "Choisir Fermer l'ancienne session et me connecter seulement si l'autre session n'est plus utile.",
                "Choisir Quitter pour conserver l'ancienne connexion.",
                "En cas de doute, demander à l'Admin de vérifier le statut dans Utilisateurs.",
            ]),
            ("Changement de mot de passe", [
                "Ouvrir la sécurité du compte.",
                "Saisir le mot de passe actuel.",
                "Saisir deux fois le nouveau mot de passe.",
                "Utiliser au moins 14 caractères pour Admin/DG et 12 pour les autres rôles.",
                "Inclure majuscule, minuscule, chiffre et symbole.",
                "Valider puis se reconnecter.",
            ]),
        ],
        "rules": [
            "Un utilisateur ne doit avoir qu'une session active.",
            "L'Admin peut déconnecter un utilisateur à distance.",
            "Ne pas laisser une session ouverte sur un appareil partagé.",
            "Ne jamais envoyer le mot de passe dans un groupe WhatsApp ou le noter à côté du PC.",
            "Après une déconnexion forcée, l'application renvoie à l'écran de connexion.",
        ],
        "errors": [
            ["Identifiant ou mot de passe incorrect", "Vérifier les majuscules, l'e-mail et le mot de passe. Demander une réinitialisation à l'Admin."],
            ["Session déjà active", "Fermer l'ancienne session ou demander à l'Admin de la couper."],
            ["Session expirée", "Se reconnecter. Une session inactive ou remplacée peut être fermée automatiquement."],
            ["Serveur inaccessible", "Vérifier Internet, le PC serveur et Cloudflare Tunnel."],
        ],
    },
    {
        "number": "02",
        "slug": "tableau-de-bord",
        "title": "Tableau de bord",
        "roles": "Tous les rôles, avec indicateurs adaptés",
        "objective": "Comprendre les indicateurs mensuels, les alertes et l'activité récente.",
        "screenshots": [("02-tableau-de-bord.png", "Vue générale des indicateurs de l'activité.")],
        "fields": [
            ["Indicateurs mensuels", "Commandes, montant reçu, production, caisse, stock et autres valeurs du mois."],
            ["Valeurs persistantes", "Dettes, commissions non payées et travailleurs non payés restent visibles jusqu'au règlement."],
            ["Alertes", "Stock faible, dettes ou autres éléments nécessitant une action."],
            ["Activité récente", "Dernières opérations, surtout pour l'Admin."],
        ],
        "procedures": [
            ("Lecture du tableau", [
                "Vérifier le mois concerné.",
                "Lire les cartes visibles pour votre rôle.",
                "Contrôler d'abord les alertes et valeurs non réglées.",
                "Cliquer sur Actualiser avant une décision importante.",
                "Ouvrir le module concerné pour consulter les détails.",
            ]),
        ],
        "rules": [
            "Les indicateurs courants repartent à zéro au changement de mois.",
            "Les impayés ne sont pas remis à zéro automatiquement.",
            "Chaque rôle ne voit que les informations nécessaires à son travail.",
            "Le tableau de bord est une synthèse : les détails restent dans les modules.",
        ],
        "errors": [
            ["Valeur ancienne", "Cliquer sur Actualiser et vérifier la connexion."],
            ["Indicateur absent", "Le rôle connecté n'a peut-être pas accès à cette information."],
            ["Montant inattendu", "Ouvrir le module concerné et contrôler les lignes du mois."],
        ],
    },
    {
        "number": "03",
        "slug": "commandes",
        "title": "Commandes",
        "roles": "Admin, Directeur Général en lecture, Gestionnaire des commandes, Caissier en lecture",
        "objective": "Enregistrer les commandes, calculer les montants, gérer dettes et avances, puis filtrer les clients.",
        "screenshots": [("06-commandes.png", "Formulaire et tableau des commandes.")],
        "fields": [
            ["Date", "Date réelle de la commande. Une date future est interdite."],
            ["Client", "Nom permettant de retrouver les dettes et avances."],
            ["Statut", "Dépositaire, Maman ou Vente cash."],
            ["Nombre de bacs", "Quantité commandée."],
            ["Montant à percevoir", "Calcul automatique selon le statut et le nombre de bacs."],
            ["Montant reçu", "Somme effectivement versée par le client."],
            ["Dette / avance", "Différence entre le montant attendu et le montant reçu."],
        ],
        "procedures": [
            ("Nouvelle commande", [
                "Ouvrir Commandes.",
                "Choisir la date du jour.",
                "Saisir le client avec une orthographe constante.",
                "Choisir Dépositaire, Maman ou Vente cash.",
                "Saisir le nombre de bacs.",
                "Vérifier le montant calculé.",
                "Saisir le montant reçu.",
                "Enregistrer puis vérifier la ligne dans le tableau.",
            ]),
            ("Utiliser une avance client", [
                "Saisir exactement le même nom de client.",
                "Consulter l'avance disponible proposée par l'application.",
                "Vérifier le montant utilisé sur la nouvelle commande.",
                "Contrôler le solde d'avance restant après enregistrement.",
            ]),
            ("Filtrer les commandes", [
                "Choisir la date ou Tout afficher.",
                "Sélectionner Tous, Mamans ou Dépositaires.",
                "Cliquer sur Actualiser.",
            ]),
        ],
        "rules": [
            "Tarif Dépositaire : 4 100 FC par bac.",
            "Tarif Maman : 6 000 FC par bac.",
            "Tarif Vente cash : 4 350 FC par bac.",
            "Un paiement insuffisant crée une dette.",
            "Un paiement supérieur réserve le solde comme avance pour la prochaine commande.",
            "Une journée clôturée ne peut plus être modifiée sans réouverture Admin.",
        ],
        "errors": [
            ["Date refusée", "La date est future ou la journée est clôturée."],
            ["Avance introuvable", "Le nom du client ne correspond pas exactement à l'ancien enregistrement."],
            ["Montant incorrect", "Vérifier le statut et le nombre de bacs."],
            ["Boutons désactivés", "Le rôle est en lecture seule."],
        ],
    },
    {
        "number": "04",
        "slug": "caisse",
        "title": "Caisse",
        "roles": "Admin, Directeur Général en lecture, Caissier",
        "objective": "Enregistrer les dépenses et dettes payées, puis contrôler les entrées et le solde.",
        "screenshots": [("08-caisse.png", "Fiche de caisse et synthèse de la journée.")],
        "fields": [
            ["Montant reçu", "Repris automatiquement depuis les commandes."],
            ["Dettes payées", "Sommes récupérées sur d'anciennes dettes."],
            ["Dépenses", "Sorties d'argent de la journée."],
            ["Détails", "Explication obligatoire des dépenses et paiements de dettes."],
            ["Solde", "Entrées moins dépenses, calculé automatiquement."],
        ],
        "procedures": [
            ("Fiche de caisse", [
                "Ouvrir Caisse et choisir la date.",
                "Contrôler le montant reçu depuis les commandes.",
                "Saisir les dettes payées aujourd'hui.",
                "Décrire les clients et montants concernés.",
                "Saisir les dépenses.",
                "Décrire chaque dépense.",
                "Vérifier le solde puis enregistrer.",
            ]),
        ],
        "rules": [
            "Ne pas saisir manuellement le montant reçu déjà calculé depuis les commandes.",
            "Toute dépense doit avoir un détail compréhensible.",
            "Les dettes payées réduisent les dettes accumulées.",
            "La date future et la journée clôturée sont bloquées.",
        ],
        "errors": [
            ["Montant reçu différent", "Contrôler les commandes de la même date."],
            ["Dette payée non reconnue", "Vérifier le client, la date et le montant encore dû."],
            ["Solde négatif", "Contrôler les dépenses et les montants saisis."],
        ],
    },
    {
        "number": "05",
        "slug": "stock",
        "title": "Stock",
        "roles": "Admin, Directeur Général en lecture, Gestionnaire de stock",
        "objective": "Gérer les paramètres, approvisionnements, sorties, journal et alertes de stock.",
        "screenshots": [
            ("04-stock.png", "Synthèse et journal du stock."),
            ("05-approvisionnement-stock.png", "Formulaire d'approvisionnement."),
        ],
        "fields": [
            ["Farine", "Quantité en sacs."],
            ["Levure", "Quantité en paquets."],
            ["Sel", "Quantité en kilogrammes."],
            ["Huile", "Quantité en litres."],
            ["Paramètres", "Seuils d'alerte et règles de consommation."],
            ["Journal", "Stock d'ouverture, entrées, sorties et clôture."],
        ],
        "procedures": [
            ("Approvisionnement", [
                "Ouvrir Stock puis Approvisionnement.",
                "Choisir la date du jour.",
                "Saisir les quantités reçues.",
                "Ajouter une observation ou référence de livraison.",
                "Enregistrer et contrôler le stock restant.",
            ]),
            ("Sortie de stock", [
                "Ouvrir Sortie de stock.",
                "Choisir la date.",
                "Saisir les quantités utilisées.",
                "Vérifier la cohérence avec la production.",
                "Enregistrer.",
            ]),
            ("Paramètres", [
                "Ouvrir Paramètres du stock.",
                "Contrôler les seuils d'alerte et équivalences.",
                "Modifier seulement avec l'accord de l'Admin.",
                "Enregistrer puis actualiser le module.",
            ]),
        ],
        "rules": [
            "Le stock ne doit pas devenir négatif.",
            "Les sacs utilisés doivent correspondre aux données de production.",
            "Une date future est interdite.",
            "Toute modification doit rester traçable dans l'historique.",
        ],
        "errors": [
            ["Stock insuffisant", "Contrôler les approvisionnements et la quantité de sortie."],
            ["Écart avec production", "Comparer les sacs utilisés dans Stock et Production."],
            ["Alerte persistante", "Le niveau reste inférieur au seuil configuré."],
        ],
    },
    {
        "number": "06",
        "slug": "production",
        "title": "Production",
        "roles": "Admin, Directeur Général en lecture, Chargé de la production, Caissier en lecture",
        "objective": "Enregistrer la production journalière et contrôler les écarts avec commandes et stock.",
        "screenshots": [("07-production.png", "Formulaire de production et calculs automatiques.")],
        "fields": [
            ["Bacs commandés", "Total attendu depuis les besoins du jour."],
            ["Livrés dépositaires", "Bacs affectés aux dépositaires."],
            ["Livrés mamans", "Bacs affectés aux mamans."],
            ["Donnés / échantillons", "Bacs non facturés."],
            ["Restants / foutus", "Production non vendue ou perdue."],
            ["Sacs utilisés", "Farine consommée pour la production."],
        ],
        "procedures": [
            ("Production journalière", [
                "Ouvrir Production.",
                "Choisir la date du jour.",
                "Saisir les bacs commandés.",
                "Répartir les bacs livrés aux dépositaires et mamans.",
                "Saisir les bacs donnés, échantillons, restants et foutus.",
                "Saisir le nombre de sacs utilisés.",
                "Lire le total produit, l'écart et le taux de couverture.",
                "Enregistrer.",
            ]),
        ],
        "rules": [
            "Le total produit est calculé automatiquement.",
            "L'écart compare production et commandes.",
            "Les sacs utilisés doivent correspondre aux sorties de stock.",
            "Le Chargé de la production ne gère pas les autres modules.",
        ],
        "errors": [
            ["Écart important", "Recompter les catégories de bacs et vérifier les commandes."],
            ["Sacs incohérents", "Comparer avec la sortie de stock du même jour."],
            ["Enregistrement refusé", "Vérifier date, clôture et rôle."],
        ],
    },
    {
        "number": "07",
        "slug": "commissions",
        "title": "Commissions",
        "roles": "Admin, Directeur Général en lecture, Gestionnaire des commandes, Caissier en lecture",
        "objective": "Consulter les commissions calculées depuis les commandes et suivre leur règlement.",
        "screenshots": [("09-commissions.png", "Synthèse et liste des commissions.")],
        "fields": [
            ["Client", "Personne ou entité concernée."],
            ["Date", "Date de la commande associée."],
            ["Bacs", "Base de calcul issue des commandes."],
            ["Commission", "Montant calculé."],
            ["Statut", "Payée ou restant à payer selon le suivi."],
        ],
        "procedures": [
            ("Consulter les commissions", [
                "Ouvrir Commissions.",
                "Choisir la date ou la période.",
                "Cliquer sur Actualiser.",
                "Contrôler le client, le nombre de bacs et le montant.",
                "Comparer avec les commandes de la même date.",
            ]),
        ],
        "rules": [
            "Les commissions proviennent des commandes enregistrées.",
            "Une commission non payée reste visible le mois suivant.",
            "Le Directeur Général et le Caissier consultent sans modifier.",
            "Toute correction doit commencer par la vérification de la commande source.",
        ],
        "errors": [
            ["Commission absente", "Vérifier que la commande source existe et que le filtre est correct."],
            ["Montant inattendu", "Contrôler statut, nombre de bacs et règle de commission."],
        ],
    },
    {
        "number": "08",
        "slug": "travailleurs-et-paies",
        "title": "Travailleurs et paies",
        "roles": "Admin, Directeur Général en lecture, Caissier",
        "objective": "Gérer les dossiers des travailleurs, l'ancienneté et les étapes de paie.",
        "screenshots": [],
        "fields": [
            ["Nom et fonction", "Identification du travailleur."],
            ["Téléphone, adresse, e-mail", "Contacts et destination des notifications."],
            ["Date d'embauche", "Base de calcul automatique de l'ancienneté."],
            ["Salaire", "Montant mensuel de référence."],
            ["Prime, avance, retenue", "Éléments qui modifient le net."],
            ["Statut de paie", "Préparée, validée ou payée."],
        ],
        "procedures": [
            ("Créer un travailleur", [
                "Ouvrir Travailleurs.",
                "Saisir le nom complet et la fonction.",
                "Renseigner téléphone, adresse et e-mail.",
                "Saisir la date d'embauche.",
                "Saisir le salaire mensuel.",
                "Choisir le statut Actif puis enregistrer.",
            ]),
            ("Enregistrer une paie", [
                "Sélectionner le travailleur.",
                "Choisir la période.",
                "Saisir brut, prime, avance et retenue.",
                "Vérifier le net calculé.",
                "Choisir le mode et le statut de paiement.",
                "Enregistrer.",
            ]),
        ],
        "rules": [
            "L'ancienneté se calcule depuis la date d'embauche.",
            "Le net ne peut pas être négatif.",
            "Les travailleurs non payés restent visibles jusqu'au paiement.",
            "Les notifications sont envoyées en arrière-plan lorsque le service e-mail est actif.",
        ],
        "errors": [
            ["Net négatif", "Les avances et retenues dépassent le brut plus la prime."],
            ["E-mail non reçu", "Vérifier l'adresse, la file d'attente et le service e-mail."],
            ["Ancienneté incorrecte", "Corriger la date d'embauche."],
        ],
    },
    {
        "number": "09",
        "slug": "rapports",
        "title": "Rapports PDF et Excel",
        "roles": "Tous les rôles, contenu limité aux modules autorisés",
        "objective": "Générer, ouvrir, imprimer et retrouver les rapports journaliers, mensuels et par période.",
        "screenshots": [
            ("10-rapport-pdf.png", "Paramètres du rapport PDF."),
            ("11-rapport-excel.png", "Paramètres du rapport Excel."),
        ],
        "fields": [
            ["Journalier", "Utilise la date de référence."],
            ["Mensuel", "Utilise le mois et l'année de la date de référence."],
            ["Période", "Utilise une date de début et une date de fin."],
            ["Caisse hebdo", "Rapport de caisse sur une semaine."],
            ["Bilan mensuel", "Synthèse de caisse du mois."],
        ],
        "procedures": [
            ("Rapport PDF", [
                "Ouvrir Rapports.",
                "Choisir le type de rapport.",
                "Renseigner la date ou la période.",
                "Cliquer sur Générer le rapport PDF.",
                "Attendre l'ouverture automatique du fichier.",
            ]),
            ("Rapport Excel", [
                "Ouvrir Rapports Excel.",
                "Choisir Journalier, Mensuel, Période, Caisse hebdo ou Bilan mensuel.",
                "Renseigner les dates nécessaires.",
                "Cliquer sur Générer le fichier Excel.",
                "Vérifier l'ouverture automatique.",
            ]),
            ("Retrouver un rapport", [
                "Cliquer sur Afficher le dossier des rapports.",
                "Rechercher le fichier par date et type.",
                "Ne pas déplacer les rapports signés liés aux clôtures sans sauvegarde.",
            ]),
        ],
        "rules": [
            "Le contenu dépend du rôle connecté.",
            "La date de fin n'est active que pour un rapport par période.",
            "Les rapports de clôture doivent rester associés à leur sauvegarde.",
            "Actualiser avant de générer si des écritures viennent d'être ajoutées.",
        ],
        "errors": [
            ["Rapport vide", "Vérifier les dates, filtres et données du module."],
            ["Fichier ne s'ouvre pas", "Utiliser Afficher le dossier des rapports."],
            ["Section absente", "Le rôle connecté n'est pas autorisé à voir cette section."],
        ],
    },
    {
        "number": "10",
        "slug": "utilisateurs",
        "title": "Utilisateurs et suivi des connexions",
        "roles": "Admin en gestion, Directeur Général en lecture",
        "objective": "Créer les comptes, attribuer les rôles, suivre les connexions et forcer une déconnexion.",
        "screenshots": [("03-utilisateurs.png", "Liste des comptes et formulaire utilisateur.")],
        "fields": [
            ["Nom complet", "Identité de la personne."],
            ["Identifiant", "Nom unique de connexion."],
            ["E-mail", "Connexion et notifications."],
            ["Rôle", "Détermine les modules et droits."],
            ["Mot de passe", "Secret initial ou nouveau mot de passe."],
            ["Connexion", "En ligne ou Hors ligne."],
            ["Plateforme / IP", "Origine de la session active."],
        ],
        "procedures": [
            ("Créer un compte", [
                "Ouvrir Utilisateurs.",
                "Saisir nom, identifiant et e-mail.",
                "Choisir le rôle.",
                "Définir un mot de passe fort.",
                "Enregistrer et vérifier la file d'e-mail.",
            ]),
            ("Modifier un compte", [
                "Charger l'utilisateur.",
                "Modifier les informations autorisées.",
                "Si le mot de passe change, avertir l'utilisateur.",
                "Enregistrer.",
            ]),
            ("Déconnecter un utilisateur", [
                "Repérer le statut En ligne.",
                "Vérifier la plateforme et la dernière activité.",
                "Cliquer sur Déconnecter.",
                "Confirmer l'action.",
            ]),
        ],
        "rules": [
            "Un seul Directeur Général peut exister.",
            "Le dernier Admin ne doit pas être supprimé.",
            "Seul l'Admin modifie les comptes.",
            "L'Admin ne doit pas couper sa propre session depuis la liste.",
            "La réinitialisation de la base exige une sauvegarde préalable.",
        ],
        "errors": [
            ["Identifiant déjà utilisé", "Choisir un identifiant unique."],
            ["Deuxième Directeur Général refusé", "Modifier ou supprimer le DG existant avant d'en créer un autre."],
            ["Mot de passe refusé", "Respecter longueur et complexité."],
            ["Utilisateur toujours en ligne", "Actualiser puis vérifier la dernière activité."],
        ],
    },
    {
        "number": "11",
        "slug": "historique-cloture-sauvegardes",
        "title": "Historique, clôture et sauvegardes",
        "roles": "Admin et Directeur Général; réouverture et sauvegarde réservées à l'Admin",
        "objective": "Auditer les actions, clôturer la journée, rouvrir si nécessaire et protéger la base.",
        "screenshots": [],
        "fields": [
            ["Historique", "50 dernières actions, affichées par pages de 10."],
            ["Filtres", "Identifiant et rôle."],
            ["Clôture", "Date à figer et statut de la journée."],
            ["Réouverture", "Motif obligatoire et action réservée à l'Admin."],
            ["Sauvegardes", "Fichier, taille, date et chemin serveur."],
        ],
        "procedures": [
            ("Consulter l'historique", [
                "Ouvrir Historique.",
                "Cliquer sur Actualiser.",
                "Filtrer par identifiant ou rôle si nécessaire.",
                "Naviguer avec Précédent et Suivant.",
                "Lire les détails de l'action.",
            ]),
            ("Clôturer la journée", [
                "Contrôler commandes, stock, production, caisse et paies.",
                "Choisir la date du jour.",
                "Cliquer sur Clôturer la journée.",
                "Vérifier le rapport et le chemin de sauvegarde.",
            ]),
            ("Rouvrir une journée", [
                "Se connecter comme Admin.",
                "Choisir la date clôturée.",
                "Saisir un motif précis.",
                "Cliquer sur Réouvrir la journée.",
                "Corriger les écritures puis clôturer à nouveau.",
            ]),
            ("Sauvegarder", [
                "Cliquer sur Sauvegarder la base.",
                "Attendre le chemin de confirmation.",
                "Vérifier le fichier dans la liste.",
                "Copier régulièrement une sauvegarde hors du PC serveur.",
            ]),
        ],
        "rules": [
            "Le Directeur Général peut clôturer mais ne peut pas rouvrir.",
            "Une clôture crée un rapport et une sauvegarde.",
            "Le motif de réouverture doit expliquer la correction.",
            "La restauration est une opération sensible réservée à l'Admin.",
            "La sauvegarde externe hebdomadaire reste obligatoire.",
        ],
        "errors": [
            ["Clôture refusée", "Vérifier le rôle, la date et les contrôles métier."],
            ["Réouverture refusée", "Seul l'Admin peut rouvrir et le motif est obligatoire."],
            ["Sauvegarde absente", "Vérifier le PC serveur et le dossier de sauvegarde."],
        ],
    },
    {
        "number": "12",
        "slug": "a-propos",
        "title": "À propos et contacts",
        "roles": "Tous les utilisateurs",
        "objective": "Identifier le responsable, les contacts et les informations de propriété de l'application.",
        "screenshots": [],
        "fields": [
            ["Responsable", "Christian Lomoto."],
            ["Entreprise", "General Investment Services (GIS)."],
            ["Téléphone", "+243 991 599 600."],
            ["E-mails", "kayboxstore@gmail.com et kayboxstore@outlook.fr."],
            ["Mentions", "Propriété, copyright et droits réservés."],
        ],
        "procedures": [
            ("Consulter les contacts", [
                "Ouvrir À propos.",
                "Lire le responsable et les coordonnées.",
                "Utiliser ces contacts pour le support ou la maintenance.",
                "Cliquer sur Actualiser si les informations ne s'affichent pas.",
            ]),
        ],
        "rules": [
            "Ne pas modifier ou supprimer les mentions légales.",
            "Ne pas diffuser les fichiers d'installation sans autorisation.",
            "Signaler les incidents avec la version de l'application et une capture d'écran.",
        ],
        "errors": [
            ["Informations absentes", "Actualiser la page ou vérifier la version installée."],
            ["Support difficile à diagnostiquer", "Communiquer la version, le module, l'heure et le message exact."],
        ],
    },
    {
        "number": "13",
        "slug": "android",
        "title": "Installation et utilisation Android",
        "roles": "Utilisateurs autorisés à travailler depuis un téléphone Android",
        "objective": "Installer l'APK, se connecter et tester l'accès distant de manière sûre.",
        "screenshots": [],
        "fields": [
            ["APK debug", "Version de recette avant signature finale."],
            ["APK release", "Version finale signée par la clé Android de l'entreprise."],
            ["Domaine", "https://app.boulangerie-lomoto.com."],
            ["Connexion", "Internet requis pour atteindre le PC serveur à distance."],
        ],
        "procedures": [
            ("Installation par câble USB", [
                "Utiliser un câble USB capable de transférer les données.",
                "Activer les options développeur et le débogage USB.",
                "Brancher et déverrouiller le téléphone.",
                "Autoriser l'ordinateur sur le téléphone.",
                "Lancer scripts/install_android_debug.ps1 depuis le projet.",
                "Ouvrir Boulangerie Lomoto après installation.",
            ]),
            ("Test distant", [
                "Désactiver le Wi-Fi du téléphone.",
                "Activer les données mobiles.",
                "Ouvrir l'APK.",
                "Se connecter et actualiser un module autorisé.",
                "Vérifier qu'une action Windows apparaît sur Android et inversement.",
            ]),
        ],
        "rules": [
            "Le téléphone doit être protégé par code, empreinte ou reconnaissance faciale.",
            "Ne pas installer l'APK depuis une source inconnue autre que General Investment Services (GIS).",
            "La clé de signature release doit être sauvegardée hors du PC serveur.",
            "Le trafic HTTP non sécurisé est désactivé dans l'APK.",
            "La sauvegarde Android automatique est désactivée pour protéger les sessions.",
        ],
        "errors": [
            ["Téléphone non détecté", "Changer de câble, activer le débogage USB et accepter l'autorisation."],
            ["Installation refusée", "Autoriser temporairement l'installation ou utiliser ADB."],
            ["Application ne charge pas", "Vérifier Internet, le domaine et Cloudflare Tunnel."],
            ["Connexion déjà active", "Fermer l'ancienne session ou demander à l'Admin."],
        ],
    },
]


def add_header_footer(document: Document, title: str) -> None:
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=3, width=Inches(6.8))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_width(table, [0.65, 4.95, 1.2])
    left, center, right = table.rows[0].cells
    if LOGO.exists():
        left.paragraphs[0].add_run().add_picture(str(LOGO), width=Cm(1.05))
    center.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(center.paragraphs[0].add_run("BOULANGERIE LOMOTO"), size=10.5, color=RED, bold=True)
    sub = center.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run(title), size=8, color=MUTED, bold=True)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right.paragraphs[0].add_run(f"v{APP_VERSION}\n"), size=8, color=NAVY, bold=True)
    set_run_font(right.paragraphs[0].add_run(date.today().strftime("%d/%m/%Y")), size=7.5, color=MUTED)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 10, 20, 10, 20)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run(f"© {date.today().year} Boulangerie Lomoto - General Investment Services (GIS) | Page "), size=8, color=MUTED)
    add_page_number(paragraph)


def add_cover(document: Document, guide: dict) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(2.8))
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(kicker.add_run(f"GUIDE MODULE {guide['number']}"), size=10, color=RED, bold=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run(guide["title"]), size=24, color=NAVY, bold=True)
    audience = document.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(audience.add_run(f"Destiné à : {guide['roles']}"), size=11, color=MUTED, italic=True)
    add_callout(document, "Objectif", guide["objective"], LIGHT_BLUE)
    document.add_page_break()


def add_checklist(document: Document, rules: list[str]) -> None:
    document.add_heading("Checklist avant de quitter le module", level=1)
    rows = [[f"Contrôle {index}", rule, "À vérifier"] for index, rule in enumerate(rules, start=1)]
    add_matrix(document, ["N°", "Contrôle", "État"], rows, [0.7, 4.8, 1.0])


def build_guide(guide: dict) -> Path:
    document = Document()
    setup_styles(document)
    add_header_footer(document, guide["title"])
    add_cover(document, guide)

    document.add_heading("1. Accès et responsabilité", level=1)
    add_matrix(
        document,
        ["Information", "Détail"],
        [
            ["Rôles concernés", guide["roles"]],
            ["Objectif du module", guide["objective"]],
            ["Version du guide", APP_VERSION],
            ["Actualisation", "Utiliser le bouton Actualiser avant de conclure qu'une donnée manque."],
        ],
        [1.8, 4.7],
    )

    for filename, caption in guide["screenshots"]:
        add_screenshot(document, filename, caption)

    document.add_heading("2. Champs et informations affichées", level=1)
    add_matrix(document, ["Champ / zone", "Utilisation"], guide["fields"], [2.05, 4.45])

    document.add_heading("3. Procédures détaillées", level=1)
    for procedure_title, steps in guide["procedures"]:
        document.add_heading(procedure_title, level=2)
        add_numbered(document, steps)

    document.add_heading("4. Règles et contrôles métier", level=1)
    add_bullets(document, guide["rules"])
    add_callout(
        document,
        "Règle commune",
        "une date future est refusée et une journée clôturée ne peut plus être modifiée sans autorisation de réouverture.",
        LIGHT_GOLD,
    )

    document.add_heading("5. Erreurs fréquentes et solutions", level=1)
    add_matrix(document, ["Problème", "Solution recommandée"], guide["errors"], [2.2, 4.3], header_fill=LIGHT_RED)

    add_checklist(document, guide["rules"])
    output = OUTPUT_DIR / f"{guide['number']}-Guide-{guide['slug']}-Boulangerie-Lomoto-{APP_VERSION}.docx"
    document.save(output)
    return output


def build_index(outputs: list[Path]) -> Path:
    document = Document()
    setup_styles(document)
    add_header_footer(document, "Index des guides par module")
    guide = {
        "number": "00",
        "title": "Index des guides utilisateur",
        "roles": "Tous les responsables et utilisateurs",
        "objective": "Identifier rapidement le guide correspondant au rôle et au module utilisé.",
    }
    add_cover(document, guide)
    document.add_heading("Guides disponibles", level=1)
    rows = []
    for item, output in zip(GUIDES, outputs):
        rows.append([item["number"], item["title"], item["roles"], output.name])
    add_matrix(document, ["N°", "Module", "Utilisateurs", "Fichier"], rows, [0.45, 1.55, 2.15, 2.35])

    document.add_heading("Répartition conseillée", level=1)
    add_matrix(
        document,
        ["Rôle", "Guides à remettre"],
        [
            ["Admin", "Tous les guides"],
            ["Directeur Général", "01, 02, 03 à 12 et 13 si utilisation Android"],
            ["Caissier", "01, 02, 03, 04, 06, 07, 08, 09, 12 et 13"],
            ["Chargé de la production", "01, 02, 06, 09, 12 et 13"],
            ["Gestionnaire de stock", "01, 02, 05, 09, 12 et 13"],
            ["Gestionnaire des commandes", "01, 02, 03, 07, 09, 12 et 13"],
        ],
        [1.55, 4.95],
    )
    output = OUTPUT_DIR / f"00-Index-guides-utilisateur-Boulangerie-Lomoto-{APP_VERSION}.docx"
    document.save(output)
    return output


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [build_guide(guide) for guide in GUIDES]
    index = build_index(outputs)
    print(index)
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
