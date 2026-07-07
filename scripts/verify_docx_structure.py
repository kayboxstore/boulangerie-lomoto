from __future__ import annotations

import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document


def main() -> None:
    if len(sys.argv) < 2:
        raise SystemExit("Usage: verify_docx_structure.py <document.docx>")
    path = Path(sys.argv[1])
    if not path.exists():
        raise SystemExit(f"Document introuvable: {path}")
    document = Document(path)
    if len(document.paragraphs) < 20:
        raise SystemExit("Document trop court.")
    if len(document.tables) < 3:
        raise SystemExit("Nombre de tableaux insuffisant.")
    empty_headers: list[tuple[int, int]] = []
    for table_index, table in enumerate(document.tables):
        if not table.rows or not table.columns:
            raise SystemExit(f"Tableau vide: {table_index}")
        for cell_index, cell in enumerate(table.rows[0].cells):
            if not cell.text.strip():
                empty_headers.append((table_index, cell_index))
    if empty_headers:
        raise SystemExit(f"En-tetes de tableaux vides: {empty_headers}")
    with ZipFile(path) as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing = sorted(required - names)
        if missing:
            raise SystemExit(f"Pieces OOXML manquantes: {missing}")
        has_media = any(name.startswith("word/media/") for name in names)
    print(
        {
            "ok": True,
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "sections": len(document.sections),
            "has_media": has_media,
            "bytes": path.stat().st_size,
        }
    )


if __name__ == "__main__":
    main()
